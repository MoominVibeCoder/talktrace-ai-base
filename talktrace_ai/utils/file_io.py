"""talktrace_ai.utils.file_io"""
from docx import Document

def docx_to_json(docx_file_path):
    doc = Document(docx_file_path)

    if not doc.tables:
        text = []

        for para in doc.paragraphs:
            cleaned = para.text.strip()
            if cleaned:  # Skip empty lines
                text.append(cleaned)

        return "\n".join(text)

    table_data = []
    table = doc.tables[0]
    headers = [table.cell(0, i).text.strip() for i in range(len(table.columns))]

    for ri, row in enumerate(table.rows):
        # Row 0 is the header row; it would otherwise emit a useless
        # {"Code": "Code", "Bezeichnung": "Bezeichnung", ...} item.
        if ri == 0:
            continue
        cell_texts = [cell.text.strip() for cell in row.cells]
        row_data = {headers[i]: cell_texts[i] for i in range(len(headers))}
        table_data.append(row_data)

    return table_data




def read_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        txt = file.read()
    return txt


def import_file(file_dict):
    if file_dict['type'] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return docx_to_json(file_dict['datapath'])
    elif file_dict['type'] == "text/plain":
        return read_txt(file_dict['datapath'])
    else:
        return None




def write_txt(file_path, text):
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(text)


def write_docx_from_text(file_path, text):
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(file_path)


