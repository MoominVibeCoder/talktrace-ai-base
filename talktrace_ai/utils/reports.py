"""talktrace_ai.utils.reports"""
import sys
import tempfile

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ._config import translate
from .qualitative import (
    CONFIDENCE_HIGH_MIN,
    CONFIDENCE_LOW_MAX,
    confidence_band_of_cell,
)
from .stats import count_teacher_impulses
from .plot_style import light_export_style

# Konfidenz-Bänder in der Codierungs-Tabelle. Eingefaerbt werden NUR die
# beiden Raender: was belastbar ist und was nachgesehen gehoeren sollte. Das
# mittlere Band ist der Normalfall (in echten Laeufen rund drei Viertel) und
# bleibt neutral — sonst ist der halbe Report bunt und die Hervorhebung
# verliert ihre Wirkung. Die Information traegt ohnehin die Prozentzahl in
# der Zelle: sie ist praeziser als jede Stufe und bleibt im SW-Druck lesbar.
# Die Schwellen selbst stehen in qualitative.py — dort, wo sie mit den
# Kalibrier-Ankern des Prompts abgeglichen sind.
CONFIDENCE_SHADES = {"high": "DDEAE6", "low": "F2DEDA"}


def _confidence_legend_parts():
    """Legenden-Segmente als [(Text, Band)] — die Renderer faerben selbst.

    Schwellen kommen aus dem Code, nie aus einem String, damit Legende und
    Einfaerbung nicht auseinanderlaufen koennen. Nur die eingefaerbten
    Baender werden erklaert; der neutrale Rest braucht keine Legende.
    """
    return [
        (f"{translate('report', 'confidence_high')} "
         f"(≥ {CONFIDENCE_HIGH_MIN} %)", "high"),
        (f"{translate('report', 'confidence_low')} "
         f"(< {CONFIDENCE_LOW_MAX + 1} %)", "low"),
    ]


def _shade_cell(cell, hex_colour):
    """Zellhintergrund setzen (python-docx kann das nicht direkt)."""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_colour)
    cell._tc.get_or_add_tcPr().append(shd)


def _shade_run(run, hex_colour):
    """Textlauf hinterlegen — faerbt die Legenden-Woerter in ihrer Bandfarbe,
    sonst erklaerte die Legende Farben, die sie selbst nicht zeigt."""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_colour)
    run._r.get_or_add_rPr().append(shd)

def remove_table_borders(table):
    tbl = table._tbl  # Access the XML element
    tblPr = tbl.tblPr

    tblBorders = tblPr.xpath('./w:tblBorders')
    if tblBorders:
        tblPr.remove(tblBorders[0])  # Remove existing borders

    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'nil')  # 'nil' removes the line
        borders.append(edge_el)

    tblPr.append(borders)

def set_row_borders(row, top=False, bottom=False, left=False, right=False, size=12, color="000000", space="0"):
    for cell in row:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        def set_border(side):
            side_el = tcBorders.find(qn(f'w:{side}'))
            if side_el is None:
                side_el = OxmlElement(f'w:{side}')
                tcBorders.append(side_el)
            side_el.set(qn('w:val'), 'single')
            side_el.set(qn('w:sz'), str(size))       # border thickness
            side_el.set(qn('w:color'), color)        # hex color
            side_el.set(qn('w:space'), space)

        if top:
            set_border('top')
        if bottom:
            set_border('bottom')
        if left:
            set_border('left')
        if right:
            set_border('right')


DEFAULT_REPORT_SECTIONS = {
    "quant": True,
    "over_time_quant": False,
    "quali": True,
    "over_time_quali": False,
    "transitions": False,
    "legend": True,
}


def _save_fig_to_png(fig, path, size=(5.5, 2.9), dpi=300):
    fig.tight_layout()
    fig.set_size_inches(*size)
    fig.savefig(path, dpi=dpi, bbox_inches="tight")


def _add_plot_to_doc(doc, fig, caption_text):
    par = doc.add_paragraph()
    par.add_run(f"{translate('report', 'figure')}: ")
    par.add_run(caption_text).italic = True
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmpfile:
        # Reports immer Light-Stil: den UI-State retten, Light anwenden, PNG
        # speichern, UI-State wiederherstellen — damit der gecachte Plot nach
        # dem Export nicht im Light-Look hängt.
        with light_export_style(fig):
            _save_fig_to_png(fig.figure, tmpfile.name)
        doc.add_picture(tmpfile.name)
    doc.add_paragraph("")


def generate_report2(
    output_path: str,
    group_name: str,
    num_pupils: int,
    num_participants: int,
    participation_rate: float,
    teacher_data: dict,
    student_data: dict,
    plot_distribution,
    num_impulses: int,
    caption: str = "",
    plot_impulse_coding=None,
    impulse_table=None,
    plot_distribution_over_time=None,
    plot_coding_over_time=None,
    dist_over_time_df=None,
    code_over_time_df=None,
    sections: dict = None,
    output_format: str = "docx",
    model_name: str = "",
    fingerprint: str = "",
    methods_text: str = "",
    plot_transitions=None,
    transitions_df=None,
    code_group_df=None,
):
    if sections is None:
        sections = dict(DEFAULT_REPORT_SECTIONS)
    else:
        sections = {**DEFAULT_REPORT_SECTIONS, **sections}

    if impulse_table is None or plot_impulse_coding is None:
        sections["quali"] = False
        sections["over_time_quali"] = False
        sections["legend"] = False

    fmt = (output_format or "docx").lower()
    if fmt == "docx":
        _build_docx_report(output_path, group_name, num_pupils, num_participants, participation_rate,
                           teacher_data, student_data, plot_distribution, num_impulses, caption,
                           plot_impulse_coding, impulse_table,
                           plot_distribution_over_time, plot_coding_over_time,
                           sections, model_name, fingerprint, methods_text,
                           plot_transitions, transitions_df, code_group_df)
    elif fmt == "pdf":
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_docx = tmp.name
        _build_docx_report(tmp_docx, group_name, num_pupils, num_participants, participation_rate,
                           teacher_data, student_data, plot_distribution, num_impulses, caption,
                           plot_impulse_coding, impulse_table,
                           plot_distribution_over_time, plot_coding_over_time,
                           sections, model_name, fingerprint, methods_text,
                           plot_transitions, transitions_df, code_group_df)
        _save_as_pdf(tmp_docx, output_path)
    elif fmt == "xlsx":
        _save_as_xlsx(output_path, group_name, num_pupils, num_participants, participation_rate,
                      teacher_data, student_data, num_impulses, impulse_table,
                      dist_over_time_df, code_over_time_df, sections, model_name, caption,
                      fingerprint, methods_text, transitions_df)
    elif fmt == "html":
        _save_as_html(output_path, group_name, num_pupils, num_participants, participation_rate,
                      teacher_data, student_data, plot_distribution, num_impulses, caption,
                      plot_impulse_coding, impulse_table,
                      plot_distribution_over_time, plot_coding_over_time,
                      sections, model_name, fingerprint, methods_text,
                      plot_transitions, transitions_df, code_group_df)
    elif fmt == "csv":
        _save_as_csv_zip(output_path, group_name, num_pupils, num_participants,
                         participation_rate, teacher_data, student_data,
                         num_impulses, impulse_table, dist_over_time_df,
                         code_over_time_df, sections, model_name, caption,
                         fingerprint, methods_text, transitions_df)
    else:
        raise ValueError(f"Unknown output_format: {output_format}")


def _build_docx_report(
    output_path, group_name, num_pupils, num_participants, participation_rate,
    teacher_data, student_data, plot_distribution, num_impulses, caption,
    plot_impulse_coding, impulse_table,
    plot_distribution_over_time, plot_coding_over_time,
    sections, model_name, fingerprint="", methods_text="",
    plot_transitions=None, transitions_df=None, code_group_df=None,
):
    doc = Document()

    # Schriften formatieren
    styles = doc.styles
    styles['Heading1'].element.rPr.rFonts.set(qn("w:asciiTheme"), "Aptos")
    styles['Heading1'].font.name = 'Aptos'
    styles['Heading1'].font.size = Pt(16)
    styles['Heading1'].font.bold = True
    styles['Heading1'].font.color.rgb = RGBColor(0, 0, 0)  # Schwarz
    styles['Heading1'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles['Heading1'].paragraph_format.space_after = Pt(0)
    styles['Heading1'].paragraph_format.space_before = Pt(0)
    styles['Heading1'].paragraph_format.line_spacing = 1

    styles['Heading2'].font.name = 'Aptos'
    styles['Heading2'].font.size = Pt(12)
    styles['Heading2'].font.bold = True
    styles['Heading2'].font.color.rgb = RGBColor(0, 0, 0)  # Schwarz
    styles['Heading2'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles['Heading2'].paragraph_format.space_after = Pt(0)
    styles['Heading2'].paragraph_format.space_before = Pt(0)
    styles['Heading2'].paragraph_format.line_spacing = 1

    styles['Normal'].font.name = 'Aptos'
    styles['Normal'].font.size = Pt(12)
    styles['Normal'].paragraph_format.space_after = Pt(0)
    styles['Normal'].paragraph_format.space_before = Pt(0)
    styles['Normal'].paragraph_format.line_spacing = 1

    # Ränder
    doc.sections[0].left_margin = Inches(0.5)
    doc.sections[0].right_margin = Inches(0.5)
    doc.sections[0].top_margin = Inches(0.5)
    doc.sections[0].bottom_margin = Inches(0.5)

    # === Titel ===
    doc.add_heading(f"{translate("report", "header")} {group_name}", level=1)

    if sections.get("quant"):
        _docx_quant_section(doc, num_pupils, num_participants, participation_rate,
                            teacher_data, student_data, plot_distribution)

    if sections.get("over_time_quant") and plot_distribution_over_time is not None:
        _add_plot_to_doc(doc, plot_distribution_over_time, translate("results", "over_time_quant_title"))

    if sections.get("quali"):
        _docx_quali_section(doc, num_impulses, plot_impulse_coding, impulse_table,
                            code_group_df)

    if sections.get("over_time_quali") and plot_coding_over_time is not None:
        _add_plot_to_doc(doc, plot_coding_over_time, translate("results", "over_time_quali_title"))

    if sections.get("transitions") and plot_transitions is not None:
        doc.add_heading(translate("results", "transitions_title"), level=2)
        doc.add_paragraph("").paragraph_format.line_spacing = 0.3
        _add_plot_to_doc(doc, plot_transitions, translate("results", "transitions_title"))
        if transitions_df is not None and not transitions_df.empty:
            par_tr = doc.add_paragraph()
            par_tr.add_run(f"{translate('report', 'table')}: ").italic = False
            par_tr.add_run(translate("results", "transitions_title")).italic = True
            codes = list(transitions_df.index)
            ncols = 1 + len(codes)
            tt = doc.add_table(rows=1 + len(codes), cols=ncols)
            tt.style = "Table Grid"
            hdr = tt.rows[0].cells
            hdr[0].text = ""
            for j, c in enumerate(codes, start=1):
                hdr[j].text = str(c)
            for i, row_code in enumerate(codes, start=1):
                cells = tt.rows[i].cells
                cells[0].text = str(row_code)
                for j, col_code in enumerate(codes, start=1):
                    v = transitions_df.iloc[i - 1, j - 1]
                    cells[j].text = f"{v * 100:.0f}%" if v else ""
            for cell in tt.rows[0].cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row in tt.rows:
                for cell in row.cells:
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.size = Pt(8)

    if sections.get("legend"):
        doc.add_paragraph("")
        par4 = doc.add_paragraph()
        par4.add_run(f"{translate('report', 'caption')}: ")
        par4.add_run(caption).italic = True
        if model_name:
            par5 = doc.add_paragraph()
            par5.add_run(f"{translate('report', 'model_used')}: ")
            par5.add_run(model_name).italic = True
        if fingerprint:
            par6 = doc.add_paragraph()
            par6.add_run(f"{translate('report', 'fingerprint')}: ")
            par6.add_run(fingerprint).italic = True
        if methods_text:
            par7 = doc.add_paragraph()
            par7.add_run(f"{translate('report', 'methods_section')}: ").bold = True
            par8 = doc.add_paragraph()
            par8.add_run(methods_text).italic = True

    doc.save(output_path)


def _docx_quant_section(doc, num_pupils, num_participants, participation_rate,
                        teacher_data, student_data, plot_distribution):
    doc.add_heading(translate("report", "section_1"), level=2)
    doc.add_paragraph("").paragraph_format.line_spacing = 0.3
    doc.add_paragraph(f"{translate("report", "class_size")}: {num_pupils}\t\t{translate("report", "pupil_count")}: {num_participants} ({translate("report", "participation_rate")}: {participation_rate:.1f}%)")
    doc.add_paragraph("").paragraph_format.line_spacing = 0.5

    par1 = doc.add_paragraph()
    par1.add_run(f"{translate("report", "table")}: ")
    par1.add_run(translate("report", "interaction_turns_teacher_pupils")).italic = True
    par1.paragraph_format.line_spacing = 1.2

    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = translate("report", "participants")
    hdr_cells[1].text = translate("report", "teacher")
#    hdr_cells[2].text = ""
 #   hdr_cells[3].text = ""
    hdr_cells[4].text = translate("report", "pupils")
  #  hdr_cells[5].text = ""

    row1 = table.rows[1].cells
    row1[0].text = ""
    row1[1].text = translate("report", "quantity")
    row1[2].text = translate("report", "length_words")
    row1[3].text = ""
    row1[4].text = translate("report", "quantity")
    row1[5].text = translate("report", "length_words")

    row2 = table.rows[2].cells
    row2[0].text = ""
    row2[1].text = "N"
    row2[2].text = "M(SD)"
    row2[3].text = ""
    row2[4].text = "N"
    row2[5].text = "M(SD)"
    
    row3 = table.rows[3].cells
    row3[0].text = translate("report", "interaction_turns")
    row3[1].text = str(teacher_data["num"])
    row3[2].text = f"{str(teacher_data["words"])} ({str(teacher_data["mean_sd"])})"
    row3[3].text = ""
    row3[4].text = str(student_data["num"])
    row3[5].text = f"{str(student_data["words"])} ({str(student_data["mean_sd"])})"
    
    # Tabelle formatieren
    for cell in row2:
        cell.paragraphs[0].runs[0].italic = True   
    hdr_cells[1].merge(hdr_cells[2])
    hdr_cells[4].merge(hdr_cells[5])
    remove_table_borders(table)
    set_row_borders(hdr_cells, top=True, bottom=True)
    set_row_borders(row2[1:3], bottom=True)
    set_row_borders(row2[4:6], bottom=True)
    set_row_borders(row3, bottom=True)

    for row in table.rows:
        for cell in row.cells[2:6]:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_plot_to_doc(doc, plot_distribution, translate("report", "distribution_of_turns"))


def _docx_code_group_table(doc, code_group_df):
    """Verteilungstabelle Code x Sprechergruppe (Lehrkraft / Schueler:innen).

    Ergaenzt den gestapelten Balkenplot um die exakten Zahlen: der Plot zeigt
    das Verhaeltnis, die Tabelle macht es zitierfaehig. Zeilensumme in einer
    Gesamt-Spalte, Spaltensumme in einer Summenzeile. Ohne Daten (oder ohne
    Sprecher-Information) faellt der Block ersatzlos weg.
    """
    if code_group_df is None or code_group_df.empty:
        return
    par = doc.add_paragraph()
    par.add_run(f"{translate('report', 'table')}: ")
    par.add_run(translate("report", "code_by_speaker")).italic = True

    group_cols = [str(c) for c in code_group_df.columns]
    headers = [translate("report", "shortcode"), *group_cols,
               translate("report", "total")]
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h

    for code, row in code_group_df.iterrows():
        cells = tbl.add_row().cells
        cells[0].text = str(code)
        for i, c in enumerate(code_group_df.columns, start=1):
            cells[i].text = str(int(row[c]))
        cells[-1].text = str(int(row.sum()))

    totals = tbl.add_row().cells
    totals[0].text = translate("report", "total")
    for i, c in enumerate(code_group_df.columns, start=1):
        totals[i].text = str(int(code_group_df[c].sum()))
    totals[-1].text = str(int(code_group_df.to_numpy().sum()))

    for cell in tbl.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    for cell in totals:
        cell.paragraphs[0].runs[0].bold = True
    for row in tbl.rows:
        for cell in row.cells[1:]:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")


def _docx_quali_section(doc, num_impulses, plot_impulse_coding, impulse_table,
                        code_group_df=None):
    doc.add_heading(translate("report", "section_2"), level=2)
    doc.add_paragraph("").paragraph_format.line_spacing = 0.3
    doc.add_paragraph(f"{translate('report', 'impulses_count')}: N = {num_impulses}")
    _add_plot_to_doc(doc, plot_impulse_coding, translate("report", "teacher_impulses"))

    _docx_code_group_table(doc, code_group_df)

    par3 = doc.add_paragraph()
    par3.add_run(f"{translate('report', 'table')}: ")
    par3.add_run(translate("report", "teacher_impulses")).italic = True

    speaker_col = translate("report", "speaker")
    statement_col = translate("report", "teacher_statement")
    has_speaker = speaker_col in impulse_table.columns

    # Code-Spalten dynamisch: Multi-Coding liefert "Code 1".."Code 3",
    # Single-Coding die klassische eine Spalte. Sprachrobust erkannt als
    # "alles, was nicht #/Sprecher/Äußerung ist".
    code_cols = [c for c in impulse_table.columns
                 if c not in ("#", speaker_col, statement_col)]
    if not code_cols:
        code_cols = [translate("report", "shortcode")]
    code_headers = [translate("report", "code")] if len(code_cols) == 1 else [str(c) for c in code_cols]

    ncols = (2 if has_speaker else 1) + 1 + len(code_cols)
    t = doc.add_table(rows=1, cols=ncols)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    col = 0
    hdr[col].text = "#"
    col += 1
    if has_speaker:
        hdr[col].text = speaker_col
        col += 1
    hdr[col].text = statement_col
    col += 1
    for h in code_headers:
        hdr[col].text = h
        col += 1

    has_confidence = False
    for i, row in impulse_table.iterrows():
        row_cells = t.add_row().cells
        col = 0
        row_cells[col].text = str(i + 1)
        col += 1
        if has_speaker:
            row_cells[col].text = str(row[speaker_col])
            col += 1
        row_cells[col].text = str(row[statement_col])
        col += 1
        for c in code_cols:
            raw = str(row[c]) if c in row.index else ""
            row_cells[col].text = raw
            # Nur die beiden Raender einfaerben; das mittlere Band und
            # handkorrigierte Zellen (ohne Wert) bleiben neutral.
            shade = CONFIDENCE_SHADES.get(confidence_band_of_cell(raw))
            if shade:
                _shade_cell(row_cells[col], shade)
                has_confidence = True
            col += 1

    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True

    for row in t.rows:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(8)

    if has_confidence:
        legend = doc.add_paragraph()
        intro = legend.add_run(f"{translate('report', 'confidence_legend')}: ")
        intro.font.size = Pt(8)
        intro.italic = True
        for i, (text, band) in enumerate(_confidence_legend_parts()):
            if i:
                sep = legend.add_run(" · ")
                sep.font.size = Pt(8)
                sep.italic = True
            run = legend.add_run(text)
            run.font.size = Pt(8)
            run.italic = True
            _shade_run(run, CONFIDENCE_SHADES[band])

    # Breiten: schmale #-/Code-Spalten, der Rest geht in die Äußerung.
    # Mit Konfidenz-Suffix ("EN (92 %)") brauchen Code-Spalten etwas Platz.
    code_width = Inches(0.9 if len(code_cols) > 1 else 0.5)
    statement_width = Inches(max(2.5, 7.3 - 0.3 - (0.8 if has_speaker else 0) - 0.9 * len(code_cols)))
    for row in t.rows:
        col = 0
        row.cells[col].width = Inches(0.3)
        col += 1
        if has_speaker:
            row.cells[col].width = Inches(0.8)
            col += 1
        row.cells[col].width = statement_width
        col += 1
        for _ in code_cols:
            row.cells[col].width = code_width
            col += 1


def _save_as_pdf(docx_path, pdf_path):
    # docx2pdf relies on Microsoft Word (Windows COM) or Pages/Word (macOS
    # AppleScript). On Linux there is no supported backend, so fail fast with
    # a dedicated marker the UI translates into a clearer message.
    if sys.platform.startswith("linux"):
        raise RuntimeError("pdf_unavailable_linux")
    try:
        from docx2pdf import convert
    except ImportError as e:
        raise RuntimeError("pdf_unavailable") from e
    try:
        convert(docx_path, pdf_path)
    except Exception as e:
        raise RuntimeError("pdf_unavailable") from e


def _safe_sheet_name(name):
    bad = set('[]:*?/\\')
    cleaned = "".join(c for c in name if c not in bad)
    return cleaned[:31] or "Sheet"


def _save_as_xlsx(output_path, group_name, num_pupils, num_participants, participation_rate,
                  teacher_data, student_data, num_impulses, impulse_table,
                  dist_over_time_df, code_over_time_df, sections, model_name, caption,
                  fingerprint="", methods_text="", transitions_df=None):
    try:
        import openpyxl  # noqa: F401
    except ImportError as e:
        raise RuntimeError("xlsx_unavailable") from e

    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        overview_rows = [
            (translate("report", "header"), group_name),
            (translate("report", "class_size"), num_pupils),
            (translate("report", "pupil_count"), num_participants),
            (translate("report", "participation_rate"), f"{participation_rate:.1f}%"),
        ]
        if sections.get("quali") and num_impulses:
            overview_rows.append((translate("report", "impulses_count"), num_impulses))
        if model_name and sections.get("legend"):
            overview_rows.append((translate("report", "model_used"), model_name))
        if caption and sections.get("legend"):
            overview_rows.append((translate("report", "caption"), caption))
        if fingerprint and sections.get("legend"):
            overview_rows.append((translate("report", "fingerprint"), fingerprint))
        if methods_text and sections.get("legend"):
            overview_rows.append((translate("report", "methods_section"), methods_text))
        pd.DataFrame(overview_rows, columns=["Key", "Value"]).to_excel(
            writer, sheet_name=_safe_sheet_name(translate("report_options", "sheet_overview")), index=False)

        if sections.get("quant"):
            quant_df = pd.DataFrame([
                {translate("report", "participants"): translate("report", "teacher"),
                 "N": teacher_data["num"],
                 translate("report", "length_words"): teacher_data["words"],
                 "M(SD)": teacher_data["mean_sd"]},
                {translate("report", "participants"): translate("report", "pupils"),
                 "N": student_data["num"],
                 translate("report", "length_words"): student_data["words"],
                 "M(SD)": student_data["mean_sd"]},
            ])
            quant_df.to_excel(writer, sheet_name=_safe_sheet_name(translate("report_options", "sheet_quant")), index=False)

        if sections.get("over_time_quant") and dist_over_time_df is not None:
            dist_over_time_df.to_excel(
                writer, sheet_name=_safe_sheet_name(translate("report_options", "sheet_quant_over_time")), index=False)

        if sections.get("quali") and impulse_table is not None:
            impulse_table.to_excel(
                writer, sheet_name=_safe_sheet_name(translate("report_options", "sheet_impulses")), index=False)

        if sections.get("over_time_quali") and code_over_time_df is not None:
            code_over_time_df.to_excel(
                writer, sheet_name=_safe_sheet_name(translate("report_options", "sheet_quali_over_time")), index=False)

        if sections.get("transitions") and transitions_df is not None and not transitions_df.empty:
            # index=True bewahrt die Code-Labels in der ersten Spalte; ohne den
            # Index waere die Matrix nicht lesbar (man wuesste nicht welcher
            # Code von welchem ausgeht).
            transitions_df.to_excel(
                writer, sheet_name=_safe_sheet_name(translate("report_options", "sheet_transitions")), index=True)


def _save_as_csv_zip(output_path, group_name, num_pupils, num_participants,
                     participation_rate, teacher_data, student_data,
                     num_impulses, impulse_table, dist_over_time_df,
                     code_over_time_df, sections, model_name, caption,
                     fingerprint="", methods_text="", transitions_df=None):
    """Long-format CSV bundle (ZIP) for R / SPSS / Stata workflows.

    Each section becomes one CSV inside the ZIP. Quantitative stats use a
    long-format ``metric/value`` layout that drops straight into ``ggplot2``,
    pandas, or ``tidyr::pivot_*`` without reshaping. ``meta.csv`` carries the
    reproducibility fingerprint and model so a single ``read.csv("meta.csv")``
    documents the run.
    """
    import zipfile
    import csv
    import io

    def _df_csv(df, index=False):
        buf = io.StringIO()
        df.to_csv(buf, index=index)
        return buf.getvalue()

    def _rows_csv(rows):
        buf = io.StringIO()
        w = csv.writer(buf)
        for r in rows:
            w.writerow(r)
        return buf.getvalue()

    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
        # --- meta.csv: provenance everyone wants on hand -------------------
        meta_rows = [("key", "value"),
                     ("group", group_name),
                     ("class_size", num_pupils),
                     ("participants", num_participants),
                     ("participation_rate_pct", f"{participation_rate:.1f}")]
        if model_name:
            meta_rows.append(("model", model_name))
        if fingerprint:
            meta_rows.append(("fingerprint", fingerprint))
        if caption:
            meta_rows.append(("legend", caption))
        if methods_text:
            meta_rows.append(("methods_text", methods_text))
        zf.writestr("meta.csv", _rows_csv(meta_rows))

        # --- quant_long.csv: one row per (speaker_role × metric) -----------
        if sections.get("quant"):
            quant_long = []
            for role, data in (("teacher", teacher_data), ("students", student_data)):
                quant_long.append({"role": role, "metric": "n_turns", "value": data.get("num")})
                quant_long.append({"role": role, "metric": "mean_words", "value": data.get("words")})
                quant_long.append({"role": role, "metric": "median_words", "value": data.get("mean_sd")})
            zf.writestr("quant_long.csv",
                        _df_csv(pd.DataFrame(quant_long, columns=["role", "metric", "value"])))

        # --- quant_over_time.csv: passes through if section enabled --------
        if sections.get("over_time_quant") and dist_over_time_df is not None:
            zf.writestr("quant_over_time.csv", _df_csv(dist_over_time_df))

        # --- impulses_long.csv: one row per coded impulse ------------------
        if sections.get("quali") and impulse_table is not None:
            zf.writestr("impulses_long.csv", _df_csv(impulse_table))

        # --- coding_over_time.csv ------------------------------------------
        if sections.get("over_time_quali") and code_over_time_df is not None:
            zf.writestr("coding_over_time.csv", _df_csv(code_over_time_df))

        # --- transitions.csv: row-stochastic Code-Übergangsmatrix ---------
        if sections.get("transitions") and transitions_df is not None and not transitions_df.empty:
            zf.writestr("transitions.csv", _df_csv(transitions_df, index=True))


def _fig_to_base64_png(fig, size=(7.5, 4.0), dpi=150):
    import base64
    import io as _io
    buf = _io.BytesIO()
    # Reports immer Light-Stil; UI-State für gecachte Plots danach restaurieren.
    with light_export_style(fig):
        fig.figure.tight_layout()
        fig.figure.set_size_inches(*size)
        fig.figure.savefig(buf, format="png", dpi=dpi, bbox_inches="tight")
    buf.seek(0)
    return base64.b64encode(buf.read()).decode("ascii")


from .text import html_escape as _html_escape


def _save_as_html(output_path, group_name, num_pupils, num_participants, participation_rate,
                  teacher_data, student_data, plot_distribution, num_impulses, caption,
                  plot_impulse_coding, impulse_table,
                  plot_distribution_over_time, plot_coding_over_time,
                  sections, model_name, fingerprint="", methods_text="",
                  plot_transitions=None, transitions_df=None, code_group_df=None):
    parts = []
    e = _html_escape
    parts.append("<!doctype html><html><head><meta charset='utf-8'>")
    parts.append(f"<title>{e(translate('report', 'header'))} {e(group_name)}</title>")
    parts.append("<style>"
                 "body{font-family:Aptos,Arial,sans-serif;max-width:900px;margin:2rem auto;padding:0 1rem;color:#222}"
                 "h1{font-size:1.6rem;text-align:center}h2{font-size:1.15rem;border-bottom:1px solid #ccc;padding-bottom:0.2rem}"
                 "table{border-collapse:collapse;width:100%;margin:0.5rem 0;font-size:0.9rem}"
                 "th,td{border:1px solid #999;padding:0.3rem 0.5rem;text-align:left;vertical-align:top}"
                 "th{background:#f2f2f2}img{max-width:100%;height:auto;display:block;margin:0.5rem 0}"
                 ".caption{font-style:italic;color:#555;font-size:0.9rem}"
                 # Konfidenz-Baender: gleiche Farben wie im DOCX, nur die
                 # beiden Raender. Die Prozentzahl in der Zelle traegt die
                 # Information auch ohne Farbe.
                 f".conf-high{{background:#{CONFIDENCE_SHADES['high']}}}"
                 f".conf-low{{background:#{CONFIDENCE_SHADES['low']}}}"
                 "</style></head><body>")
    parts.append(f"<h1>{e(translate('report', 'header'))} {e(group_name)}</h1>")

    if sections.get("quant"):
        parts.append(f"<h2>{e(translate('report', 'section_1'))}</h2>")
        parts.append(f"<p>{e(translate('report', 'class_size'))}: {e(num_pupils)} &nbsp;&nbsp; "
                     f"{e(translate('report', 'pupil_count'))}: {e(num_participants)} "
                     f"({e(translate('report', 'participation_rate'))}: {participation_rate:.1f}%)</p>")
        parts.append("<table><thead><tr>"
                     f"<th>{e(translate('report', 'participants'))}</th>"
                     f"<th>N</th><th>{e(translate('report', 'length_words'))} M(SD)</th>"
                     "</tr></thead><tbody>")
        parts.append(f"<tr><td>{e(translate('report', 'teacher'))}</td>"
                     f"<td>{e(teacher_data['num'])}</td>"
                     f"<td>{e(teacher_data['words'])} ({e(teacher_data['mean_sd'])})</td></tr>")
        parts.append(f"<tr><td>{e(translate('report', 'pupils'))}</td>"
                     f"<td>{e(student_data['num'])}</td>"
                     f"<td>{e(student_data['words'])} ({e(student_data['mean_sd'])})</td></tr>")
        parts.append("</tbody></table>")
        if plot_distribution is not None:
            b64 = _fig_to_base64_png(plot_distribution)
            parts.append(f"<p class='caption'>{e(translate('report', 'figure'))}: {e(translate('report', 'distribution_of_turns'))}</p>")
            parts.append(f"<img src='data:image/png;base64,{b64}' alt='distribution'>")

    if sections.get("over_time_quant") and plot_distribution_over_time is not None:
        b64 = _fig_to_base64_png(plot_distribution_over_time)
        parts.append(f"<h2>{e(translate('results', 'over_time_quant_title'))}</h2>")
        parts.append(f"<img src='data:image/png;base64,{b64}' alt='participation over time'>")

    if sections.get("quali") and impulse_table is not None and plot_impulse_coding is not None:
        parts.append(f"<h2>{e(translate('report', 'section_2'))}</h2>")
        parts.append(f"<p>{e(translate('report', 'impulses_count'))}: N = {e(num_impulses)}</p>")
        b64 = _fig_to_base64_png(plot_impulse_coding)
        parts.append(f"<img src='data:image/png;base64,{b64}' alt='impulses coding'>")

        # Verteilungstabelle Code x Sprechergruppe — gleiche Zahlen wie der
        # gestapelte Balkenplot, nur zitierfaehig (siehe _docx_code_group_table).
        if code_group_df is not None and not code_group_df.empty:
            parts.append(f"<p class='caption'>{e(translate('report', 'table'))}: "
                         f"{e(translate('report', 'code_by_speaker'))}</p>")
            parts.append(f"<table><thead><tr><th>{e(translate('report', 'shortcode'))}</th>")
            for c in code_group_df.columns:
                parts.append(f"<th>{e(c)}</th>")
            parts.append(f"<th>{e(translate('report', 'total'))}</th></tr></thead><tbody>")
            for code, row in code_group_df.iterrows():
                parts.append(f"<tr><td>{e(code)}</td>")
                for c in code_group_df.columns:
                    parts.append(f"<td>{int(row[c])}</td>")
                parts.append(f"<td>{int(row.sum())}</td></tr>")
            parts.append(f"<tr><th>{e(translate('report', 'total'))}</th>")
            for c in code_group_df.columns:
                parts.append(f"<th>{int(code_group_df[c].sum())}</th>")
            parts.append(f"<th>{int(code_group_df.to_numpy().sum())}</th></tr>")
            parts.append("</tbody></table>")

        speaker_col = translate("report", "speaker")
        statement_col = translate("report", "teacher_statement")
        has_speaker = speaker_col in impulse_table.columns
        # Code-Spalten dynamisch (Multi-Coding: "Code 1".."Code 3";
        # Single-Coding: eine Spalte) — gleiches Erkennungsmuster wie im
        # DOCX-Builder.
        code_cols = [c for c in impulse_table.columns
                     if c not in ("#", speaker_col, statement_col)]
        if not code_cols:
            code_cols = [translate("report", "shortcode")]
        code_headers = [translate("report", "code")] if len(code_cols) == 1 else [str(c) for c in code_cols]
        parts.append("<table><thead><tr><th>#</th>")
        if has_speaker:
            parts.append(f"<th>{e(speaker_col)}</th>")
        parts.append(f"<th>{e(statement_col)}</th>")
        for h in code_headers:
            parts.append(f"<th>{e(h)}</th>")
        parts.append("</tr></thead><tbody>")
        has_confidence = False
        for i, row in impulse_table.iterrows():
            parts.append(f"<tr><td>{i + 1}</td>")
            if has_speaker:
                parts.append(f"<td>{e(row[speaker_col])}</td>")
            parts.append(f"<td>{e(row[statement_col])}</td>")
            for c in code_cols:
                raw = str(row[c]) if c in row.index else ""
                band = confidence_band_of_cell(raw)
                cls = f" class='conf-{band}'" if band in CONFIDENCE_SHADES else ""
                if cls:
                    has_confidence = True
                parts.append(f"<td{cls}>{e(raw)}</td>")
            parts.append("</tr>")
        parts.append("</tbody></table>")
        if has_confidence:
            seg = " · ".join(
                f"<span class='conf-{band}'>{e(text)}</span>"
                for text, band in _confidence_legend_parts()
            )
            parts.append(f"<p class='caption'>"
                         f"{e(translate('report', 'confidence_legend'))}: {seg}</p>")

    if sections.get("over_time_quali") and plot_coding_over_time is not None:
        b64 = _fig_to_base64_png(plot_coding_over_time)
        parts.append(f"<h2>{e(translate('results', 'over_time_quali_title'))}</h2>")
        parts.append(f"<img src='data:image/png;base64,{b64}' alt='codes over time'>")

    if sections.get("transitions") and plot_transitions is not None:
        parts.append(f"<h2>{e(translate('results', 'transitions_title'))}</h2>")
        b64 = _fig_to_base64_png(plot_transitions, size=(5.5, 5.0))
        parts.append(f"<img src='data:image/png;base64,{b64}' alt='code transitions'>")
        if transitions_df is not None and not transitions_df.empty:
            codes = list(transitions_df.index)
            parts.append("<table><thead><tr><th></th>")
            for c in codes:
                parts.append(f"<th>{e(c)}</th>")
            parts.append("</tr></thead><tbody>")
            for i, row_code in enumerate(codes):
                parts.append(f"<tr><th>{e(row_code)}</th>")
                for j, _col_code in enumerate(codes):
                    v = transitions_df.iloc[i, j]
                    cell = f"{v * 100:.0f}%" if v else ""
                    parts.append(f"<td style='text-align:right'>{cell}</td>")
                parts.append("</tr>")
            parts.append("</tbody></table>")

    if sections.get("legend"):
        if caption:
            parts.append(f"<p class='caption'>{e(translate('report', 'caption'))}: {e(caption)}</p>")
        if model_name:
            parts.append(f"<p class='caption'>{e(translate('report', 'model_used'))}: {e(model_name)}</p>")
        if fingerprint:
            parts.append(f"<p class='caption'>{e(translate('report', 'fingerprint'))}: <code>{e(fingerprint)}</code></p>")
        if methods_text:
            parts.append(f"<p class='caption'><strong>{e(translate('report', 'methods_section'))}:</strong></p>")
            parts.append(f"<p class='caption' style='white-space:pre-wrap'>{e(methods_text)}</p>")

    parts.append("</body></html>")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("".join(parts))