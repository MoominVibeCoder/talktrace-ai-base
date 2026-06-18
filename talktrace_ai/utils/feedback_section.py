"""Pure logic for the Feedback tab: prompt assembly + DOCX rendering.

Analog to ``methods_section.py`` (scholarly prose from analysis numbers), but
the prose is produced by an LLM rather than a template. This module stays
Shiny-free and side-effect-free so it is unit-testable:

  * ``build_metrics`` / ``extract_code_definitions`` / ``teacher_code_profile``
    distil the analysis surface into compact, model-readable blocks.
  * ``build_feedback_prompts`` assembles the (system, user) prompt pair with a
    FIXED, curated reference list — the model is told to cite only from it, so
    it cannot fabricate sources.
  * ``write_feedback_docx`` renders the (possibly user-edited) feedback text to
    a Word document.

The feedback covers the TEACHER's dialogic moves (T-SEDA codes the teacher's
turns) plus the quantitative talk metrics, and is structured as
Strengths / Areas for development / Concrete implementation tips / Sources.
"""
from __future__ import annotations

import math
import re

from .codebook_hierarchy import _extract_code


# Fixed, curated reference list. Appended deterministically and reproduced
# verbatim by the model under "Sources" — never invented. Real, well-known
# works on dialogic teaching + the T-SEDA toolkit the codebook is based on.
REFERENCES = [
    "Alexander, R. J. (2020). A Dialogic Teaching Companion. London: Routledge.",
    "Mercer, N., & Littleton, K. (2007). Dialogue and the Development of "
    "Children's Thinking: A Sociocultural Approach. London: Routledge.",
    "Cambridge Educational Dialogue Research (CEDiR) Group (2023). T-SEDA: "
    "Teacher Scheme for Educational Dialogue Analysis (v9). University of Cambridge.",
]

# Localized section headings. Single source for both the prompt instruction and
# the DOCX heading detection, so the two never drift.
FEEDBACK_HEADINGS = {
    "de": {
        "strengths": "Stärken",
        "development": "Entwicklungsfelder",
        "tips": "Konkrete Umsetzungstipps",
        "sources": "Quellen",
    },
    "en": {
        "strengths": "Strengths",
        "development": "Areas for development",
        "tips": "Concrete implementation tips",
        "sources": "Sources",
    },
}


# ---------------------------------------------------------------------------
# Defensive coercion helpers
# ---------------------------------------------------------------------------

def _num(value, default=0):
    """Coerce to int when whole, else float; fall back to ``default``."""
    try:
        if value is None:
            return default
        f = float(value)
        if not math.isfinite(f):  # inf / nan -> documented fallback
            return default
        return int(f) if f == int(f) else round(f, 1)
    except (TypeError, ValueError):
        return default


def _row_val(stats_df, speaker, col, default=0):
    """Read one cell from a dialog_stats DataFrame, defensively."""
    if stats_df is None or getattr(stats_df, "empty", True):
        return default
    try:
        m = stats_df.loc[stats_df["Sprecher"] == speaker, col]
        return m.values[0] if not m.empty else default
    except Exception:
        return default


# ---------------------------------------------------------------------------
# Markdown cleanup
# ---------------------------------------------------------------------------

def clean_markdown(text: str) -> str:
    """Strip the Markdown syntax models tend to emit, leaving readable plain
    text for the editable field and the export.

    The prompt already asks for plain text; this is the belt-and-suspenders
    pass that removes any residue: ATX ``#`` headings → plain heading lines,
    ``* ``/``+ `` bullets → ``- ``, blockquote ``> `` markers, and
    ``**bold**`` / ``*italic*`` / ``__bold__`` emphasis markers. Heading
    detection in :func:`write_feedback_docx` still works on the plain result.
    """
    out = []
    for raw in (text or "").splitlines():
        line = re.sub(r"^\s{0,3}#{1,6}\s*", "", raw)        # ATX headings
        line = re.sub(r"^(\s*)[*+]\s+", r"\1- ", line)       # * / + bullets -> -
        line = re.sub(r"^\s{0,3}>\s?", "", line)              # blockquote
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)          # **bold**
        line = re.sub(r"__(.+?)__", r"\1", line)              # __bold__
        line = re.sub(r"\*(\S(?:.*?\S)?)\*", r"\1", line)     # *italic*
        line = line.replace("**", "")                          # stray markers
        out.append(line)
    return "\n".join(out).strip()


# ---------------------------------------------------------------------------
# Analysis-surface distillation
# ---------------------------------------------------------------------------

def build_metrics(stats_df, *, num_participants=None, participation_rate=None,
                  num_pupils=None, teacher_name="LEHRER",
                  students_label="Schüler:innen") -> dict:
    """Quantitative talk metrics from a dialog_stats frame + scalars.

    ``stats_df`` has one row for the teacher (Sprecher == teacher_name) and one
    aggregate "Schüler:innen" row, with columns Anzahl_Beitraege /
    Gesamt_Woerter / Durchschnitt_Woerter / Median_Woerter (the German column
    names are internal keys). All values degrade to 0 when absent.
    """
    t_turns = _num(_row_val(stats_df, teacher_name, "Anzahl_Beitraege"))
    t_words = _num(_row_val(stats_df, teacher_name, "Gesamt_Woerter"))
    t_avg = _num(_row_val(stats_df, teacher_name, "Durchschnitt_Woerter"))
    t_med = _num(_row_val(stats_df, teacher_name, "Median_Woerter"))
    s_turns = _num(_row_val(stats_df, students_label, "Anzahl_Beitraege"))
    s_words = _num(_row_val(stats_df, students_label, "Gesamt_Woerter"))
    s_avg = _num(_row_val(stats_df, students_label, "Durchschnitt_Woerter"))
    s_med = _num(_row_val(stats_df, students_label, "Median_Woerter"))
    total_words = (t_words or 0) + (s_words or 0)
    share = round((t_words / total_words * 100), 1) if total_words else 0
    return {
        "teacher_turns": t_turns,
        "teacher_words": t_words,
        "teacher_avg_words": t_avg,
        "teacher_median_words": t_med,
        "student_turns": s_turns,
        "student_words": s_words,
        "student_avg_words": s_avg,
        "student_median_words": s_med,
        "total_words": total_words,
        "teacher_talk_share": share,
        "num_participants": _num(num_participants),
        "num_pupils": _num(num_pupils),
        "participation_rate": _num(participation_rate),
    }


def _label_desc(entry: dict):
    """Pull (label, description) from a codebook entry (DE or EN keys)."""
    label = (entry.get("Bezeichnung") or entry.get("Label")
             or entry.get("bezeichnung") or entry.get("label") or "")
    desc = (entry.get("Beschreibung") or entry.get("Description")
            or entry.get("beschreibung") or entry.get("description") or "")
    return str(label).strip(), str(desc).strip()


def extract_code_definitions(codebook) -> list:
    """List of (code, label, description) tuples from the codebook payload."""
    out = []
    if not isinstance(codebook, list):
        return out
    seen = set()
    for entry in codebook:
        if not isinstance(entry, dict):
            continue
        code = _extract_code(entry)
        if not code or code in seen:
            continue
        seen.add(code)
        label, desc = _label_desc(entry)
        out.append((code, label, desc))
    return out


# Speaker labels that count as "the teacher" when matching coded rows, beyond
# the configured teacher name (LLM output is not always normalized).
_TEACHER_ALIASES = {"lehrperson", "lehrer", "lehrerin", "lehrkraft", "teacher"}


def teacher_code_profile(analysis_df, teacher_name="LEHRER") -> dict:
    """Frequency of T-SEDA codes on the teacher's turns.

    Counts the Shortcode column for rows whose Sprecher matches the teacher
    (case-insensitive, plus common aliases). Falls back to all coded rows if
    nothing matches the teacher (unusual labels) so the profile is never empty
    when codings exist.
    """
    profile: dict = {}
    if analysis_df is None or getattr(analysis_df, "empty", True):
        return profile
    tl = (teacher_name or "").strip().lower()

    def _tally(only_teacher):
        prof: dict = {}
        for _, row in analysis_df.iterrows():
            code = str(row.get("Shortcode", "")).strip()
            if not code:
                continue
            if only_teacher:
                spk = str(row.get("Sprecher", "")).strip().lower()
                if spk != tl and spk not in _TEACHER_ALIASES:
                    continue
            prof[code] = prof.get(code, 0) + 1
        return prof

    profile = _tally(only_teacher=True)
    if not profile:
        profile = _tally(only_teacher=False)
    return profile


# ---------------------------------------------------------------------------
# Prompt assembly
# ---------------------------------------------------------------------------

def _metrics_block(lang: str, m: dict) -> str:
    if lang == "de":
        return (
            f"- Redeanteil der Lehrkraft: {m['teacher_talk_share']}% der Wörter "
            f"({m['teacher_words']} von {m['total_words']} Wörtern)\n"
            f"- Lehrkraft: {m['teacher_turns']} Redebeiträge "
            f"(Ø {m['teacher_avg_words']} Wörter, Median {m['teacher_median_words']})\n"
            f"- Schüler:innen: {m['student_turns']} Redebeiträge "
            f"(Ø {m['student_avg_words']} Wörter, Median {m['student_median_words']})\n"
            f"- Aktiv beteiligte Lernende: {m['num_participants']} von {m['num_pupils']}\n"
            f"- Beteiligungsquote: {m['participation_rate']}%"
        )
    return (
        f"- Teacher talk-share: {m['teacher_talk_share']}% of words "
        f"({m['teacher_words']} of {m['total_words']} words)\n"
        f"- Teacher: {m['teacher_turns']} turns "
        f"(avg {m['teacher_avg_words']} words, median {m['teacher_median_words']})\n"
        f"- Students: {m['student_turns']} turns "
        f"(avg {m['student_avg_words']} words, median {m['student_median_words']})\n"
        f"- Students who spoke: {m['num_participants']} of {m['num_pupils']}\n"
        f"- Participation rate: {m['participation_rate']}%"
    )


def _codes_block(code_definitions) -> str:
    if not code_definitions:
        return "(—)"
    lines = []
    for code, label, desc in code_definitions:
        if label and desc:
            lines.append(f"- {code} ({label}): {desc}")
        elif label:
            lines.append(f"- {code}: {label}")
        elif desc:
            lines.append(f"- {code}: {desc}")
        else:
            lines.append(f"- {code}")
    return "\n".join(lines)


def _profile_block(lang: str, code_definitions, code_profile) -> str:
    code_profile = code_profile or {}
    if not code_definitions and not code_profile:
        return "(—)"
    # Iterate codebook order when available so absent codes (count 0) are shown
    # too; this is what makes "which moves are missing" legible to the model.
    order = [c for c, _, _ in code_definitions] if code_definitions else list(code_profile)
    for c in code_profile:
        if c not in order:
            order.append(c)
    label_of = {c: lbl for c, lbl, _ in (code_definitions or [])}
    times = "×" if lang == "de" else "x"
    lines = []
    for c in order:
        n = code_profile.get(c, 0)
        lbl = label_of.get(c, "")
        lines.append(f"- {c} ({lbl}): {n}{times}" if lbl else f"- {c}: {n}{times}")
    return "\n".join(lines)


def _references_block() -> str:
    return "\n".join(f"- {r}" for r in REFERENCES)


_SYSTEM_DE = (
    "Du bist eine erfahrene Fachleitung und gibst einer Lehrkraft ein "
    "formatives, wertschätzendes Feedback zu EINER analysierten Unterrichtsstunde "
    "(dialogisches Unterrichtsgespräch). Dein Feedback stützt sich auf die "
    "Forschung zum dialogischen Unterricht (Alexander; Mercer & Littleton) und "
    "auf das Cambridge-T-SEDA-Raster.\n\n"
    "Regeln:\n"
    "- Wende dich direkt an die Lehrkraft (Anrede: Sie). Formativ und "
    "konstruktiv, niemals benotend — eine Reflexionshilfe, keine Bewertung.\n"
    "- Stütze JEDE Aussage ausschließlich auf die gelieferten Kennzahlen und das "
    "Code-Profil. Erfinde keine Zahlen.\n"
    "- Zitiere AUSSCHLIESSLICH aus der gelieferten Quellenliste; erfinde keine "
    "Autoren, Jahre oder DOIs.\n"
    "- Antworte auf Deutsch, ca. 400–600 Wörter.\n"
    "- Gliedere den Text in GENAU vier Abschnitte. Setze jede Abschnitts-"
    "überschrift als eigene Zeile in Klartext: {strengths}, dann {development}, "
    "dann {tips}, dann {sources}.\n"
    "- Verwende KEINE Markdown-Formatierung (keine #, *, ** oder _). Für "
    "Aufzählungen genügt ein einfacher Spiegelstrich (- ) am Zeilenanfang.\n"
    "- Die Tipps müssen konkret und bereits in der nächsten Stunde umsetzbar "
    "sein (z. B. konkrete Frageformate, Gesprächsroutinen, Wartezeit).\n"
    "- Liste unter '{sources}' die gelieferte Quellenliste unverändert auf."
)

_SYSTEM_EN = (
    "You are an experienced teacher educator giving a teacher formative, "
    "supportive feedback on ONE analysed lesson (a dialogic classroom "
    "discussion). Your feedback is grounded in research on dialogic teaching "
    "(Alexander; Mercer & Littleton) and the Cambridge T-SEDA framework.\n\n"
    "Rules:\n"
    "- Address the teacher directly (second person). Formative and "
    "constructive, never grading — an aid for reflection, not an assessment.\n"
    "- Base EVERY claim only on the supplied metrics and code profile. Do not "
    "invent numbers.\n"
    "- Cite ONLY from the supplied reference list; never fabricate authors, "
    "years or DOIs.\n"
    "- Answer in English, about 400–600 words.\n"
    "- Structure the text into EXACTLY four sections. Put each section heading "
    "on its own line as plain text: {strengths}, then {development}, then "
    "{tips}, then {sources}.\n"
    "- Do NOT use any Markdown formatting (no #, *, ** or _). A leading hyphen "
    "(- ) is fine for bullet points.\n"
    "- The tips must be concrete and usable in the very next lesson (e.g. "
    "specific question formats, talk routines, wait time).\n"
    "- Under '{sources}', reproduce the supplied reference list verbatim."
)

_USER_DE = (
    "Hier sind die Auswertungsdaten einer Unterrichtsstunde.\n\n"
    "QUANTITATIVE KENNZAHLEN:\n{metrics}\n\n"
    "T-SEDA-CODE-DEFINITIONEN (die Codes erfassen die dialogischen Moves der "
    "Lehrkraft):\n{codes}\n\n"
    "CODE-PROFIL DER LEHRKRAFT (wie oft welcher Move in dieser Stunde "
    "vorkam):\n{profile}\n\n"
    "QUELLEN (nur diese zitieren, am Ende unter '{sources}' unverändert "
    "auflisten):\n{references}\n\n"
    "Schreibe jetzt das formative Feedback in den vier vorgegebenen Abschnitten."
)

_USER_EN = (
    "Here is the analysis data for one lesson.\n\n"
    "QUANTITATIVE METRICS:\n{metrics}\n\n"
    "T-SEDA CODE DEFINITIONS (the codes capture the teacher's dialogic "
    "moves):\n{codes}\n\n"
    "TEACHER CODE PROFILE (how often each move occurred in this "
    "lesson):\n{profile}\n\n"
    "SOURCES (cite only these; reproduce them verbatim under '{sources}' at the "
    "end):\n{references}\n\n"
    "Now write the formative feedback in the four prescribed sections."
)


def build_feedback_prompts(*, lang="de", model="", metrics=None,
                           code_definitions=None, code_profile=None):
    """Assemble the (system_prompt, user_prompt) pair for the feedback call.

    Pure: takes the already-distilled metrics dict, code definitions and code
    profile and returns two strings. ``model`` is accepted for parity with
    other builders but is not injected into the prose.
    """
    lang = "de" if lang not in ("de", "en") else lang
    metrics = metrics or {}
    h = FEEDBACK_HEADINGS[lang]

    system_tpl = _SYSTEM_DE if lang == "de" else _SYSTEM_EN
    user_tpl = _USER_DE if lang == "de" else _USER_EN

    system_prompt = system_tpl.format(
        strengths=h["strengths"], development=h["development"],
        tips=h["tips"], sources=h["sources"],
    )
    user_prompt = user_tpl.format(
        metrics=_metrics_block(lang, {**_metric_defaults(), **metrics}),
        codes=_codes_block(code_definitions),
        profile=_profile_block(lang, code_definitions, code_profile),
        references=_references_block(),
        sources=h["sources"],
    )
    return system_prompt, user_prompt


def _metric_defaults() -> dict:
    return {
        "teacher_turns": 0, "teacher_words": 0, "teacher_avg_words": 0,
        "teacher_median_words": 0, "student_turns": 0, "student_words": 0,
        "student_avg_words": 0, "student_median_words": 0, "total_words": 0,
        "teacher_talk_share": 0, "num_participants": 0, "num_pupils": 0,
        "participation_rate": 0,
    }


# ---------------------------------------------------------------------------
# DOCX rendering of the (possibly edited) feedback text
# ---------------------------------------------------------------------------

def _strip_inline_md(s: str) -> str:
    """Drop the few markdown markers the model may emit inline."""
    return s.replace("**", "").replace("__", "").strip()


def _heading_match(line: str, headings) -> bool:
    """True if ``line`` is one of the four section headings (md-tolerant)."""
    cleaned = line.lstrip("#* ").rstrip(": ").strip().lower()
    cleaned = _strip_inline_md(cleaned).lower()
    return cleaned in {v.lower() for v in headings.values()}


def write_feedback_docx(path, text, *, lang="de", doc_title="", disclaimer="") -> None:
    """Render the feedback text to a Word document at ``path``.

    Detects the four section headings (and markdown ``#``/``**`` variants) and
    renders them as bold headings; bullet lines as a bulleted list; everything
    else as justified body paragraphs. Adds a small italic disclaimer footer.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lang = "de" if lang not in ("de", "en") else lang
    headings = FEEDBACK_HEADINGS[lang]
    text = clean_markdown(text)  # tolerate Markdown pasted into the edit field

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    if doc_title:
        h = doc.add_heading(level=0)
        run = h.add_run(doc_title)
        run.font.color.rgb = RGBColor(0, 0, 0)

    for raw in (text or "").splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        if _heading_match(stripped, headings):
            label = _strip_inline_md(stripped.lstrip("#* ").rstrip(": ").strip())
            hp = doc.add_heading(level=2)
            run = hp.add_run(label)
            run.font.color.rgb = RGBColor(0, 0, 0)
            continue
        if stripped[:2] in ("- ", "* ", "• "):
            doc.add_paragraph(_strip_inline_md(stripped[2:]), style="List Bullet")
            continue
        p = doc.add_paragraph(_strip_inline_md(stripped))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if disclaimer:
        doc.add_paragraph()
        foot = doc.add_paragraph()
        run = foot.add_run(disclaimer)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(path)
