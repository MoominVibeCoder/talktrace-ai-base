"""Intercoder – export (xlsx / csv-zip / json / html / docx / pdf)."""
import json
import os
import sys
import tempfile

import pandas as pd

from .text import html_escape as _html_escape


def _testing_overview_rows(result):
    k = result.get("kappa", float("nan"))
    pa = result.get("percent_agreement", float("nan"))
    alpha = result.get("krippendorff_alpha", float("nan"))
    gwet = result.get("gwet_ac1", float("nan"))
    bp = result.get("brennan_prediger", float("nan"))
    ci_low = result.get("kappa_ci_low")
    ci_high = result.get("kappa_ci_high")
    return [
        ("Cohen's κ", f"{k:.3f}" if k == k else "n/a"),
        ("Percent agreement", f"{pa*100:.1f} %" if pa == pa else "n/a"),
        ("Krippendorff's α", f"{alpha:.3f}" if alpha == alpha else "n/a"),
        ("Gwet's AC1", f"{gwet:.3f}" if gwet == gwet else "n/a"),
        ("Brennan-Prediger κ", f"{bp:.3f}" if bp == bp else "n/a"),
        ("Confidence interval",
         f"[{ci_low:.3f}, {ci_high:.3f}]" if ci_low is not None and ci_high is not None else "n/a"),
        ("Aligned pairs", result.get("n_pairs", "")),
        ("Coded by both", result.get("n_both", "")),
        ("Only in A", result.get("n_only_a", "")),
        ("Only in B", result.get("n_only_b", "")),
    ]


def export_testing_agreement(output_path, result, sheet_overview="Overview", sheet_confusion="Confusion", sheet_per_code="Per-Code", sheet_pairs="Pairs"):
    """Export intercoder-agreement results as a multi-sheet XLSX.

    result: dict returned by compute_intercoder_agreement().
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font
        import openpyxl.utils.dataframe
    except ImportError as e:
        raise RuntimeError("xlsx_unavailable") from e

    wb = Workbook()

    # --- Sheet 1: Overview --------------------------------------------------
    ws = wb.active
    ws.title = sheet_overview
    k = result.get("kappa", float("nan"))
    pa = result.get("percent_agreement", float("nan"))
    alpha = result.get("krippendorff_alpha", float("nan"))
    gwet = result.get("gwet_ac1", float("nan"))
    bp = result.get("brennan_prediger", float("nan"))
    rows = [
        ["Metric", "Value"],
        ["Cohen's κ", f"{k:.3f}" if k == k else "n/a"],
        ["Percent agreement", f"{pa*100:.1f} %" if pa == pa else "n/a"],
        ["Krippendorff's α", f"{alpha:.3f}" if alpha == alpha else "n/a"],
        ["Gwet's AC1", f"{gwet:.3f}" if gwet == gwet else "n/a"],
        ["Brennan-Prediger κ", f"{bp:.3f}" if bp == bp else "n/a"],
        ["Confidence interval", f"[{result.get('kappa_ci_low', float('nan')):.3f}, {result.get('kappa_ci_high', float('nan')):.3f}]" if result.get("kappa_ci_low") is not None else "n/a"],
        ["", ""],
        ["Aligned pairs", result.get("n_pairs", "")],
        ["Coded by both", result.get("n_both", "")],
        ["Only in A", result.get("n_only_a", "")],
        ["Only in B", result.get("n_only_b", "")],
    ]
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row, 1):
            c = ws.cell(row=i, column=j, value=val)
            if i == 1:
                c.font = Font(bold=True)

    # --- Sheet 2: Confusion Matrix ------------------------------------------
    ws2 = wb.create_sheet(title=sheet_confusion)
    cm = result.get("confusion")
    if cm is not None and not cm.empty:
        openpyxl.utils.dataframe.dataframe_to_rows(cm, index=True, header=True)
        # openpyxl returns rows generator; iterate manually
        for i, row in enumerate(openpyxl.utils.dataframe.dataframe_to_rows(cm, index=True, header=True), 1):
            for j, val in enumerate(row, 1):
                c = ws2.cell(row=i, column=j, value=val)
                if i == 1:
                    c.font = Font(bold=True)
        ws2.cell(row=1, column=1, value="A \\ B")

    # --- Sheet 3: Per-Code metrics ------------------------------------------
    ws3 = wb.create_sheet(title=sheet_per_code)
    per_code = result.get("per_code")
    if per_code is not None and not per_code.empty:
        for i, row in enumerate(openpyxl.utils.dataframe.dataframe_to_rows(per_code, index=False, header=True), 1):
            for j, val in enumerate(row, 1):
                c = ws3.cell(row=i, column=j, value=val)
                if i == 1:
                    c.font = Font(bold=True)

    # --- Sheet 4: Pairs (Impuls, Code A, Code B) ----------------------------
    ws4 = wb.create_sheet(title=sheet_pairs)
    pairs = result.get("pairs")
    if pairs is not None:
        hdr = ["Impuls", "Code A", "Code B"]
        for j, val in enumerate(hdr, 1):
            ws4.cell(row=1, column=j, value=val).font = Font(bold=True)
        for i, item in enumerate(pairs, 2):
            ws4.cell(row=i, column=1, value=item.get("impuls"))
            ws4.cell(row=i, column=2, value=item.get("code_a"))
            ws4.cell(row=i, column=3, value=item.get("code_b"))

    wb.save(output_path)


def export_testing_agreement_csv_zip(output_path, result, labels=None):
    """Bundle four CSVs (overview, confusion, per_code, pairs) into a ZIP."""
    import zipfile
    import csv
    import io

    labels = labels or {}
    name_overview = (labels.get("sheet_overview") or "Overview")
    name_confusion = (labels.get("sheet_confusion") or "Confusion")
    name_per_code = (labels.get("sheet_per_code") or "Per-Code")
    name_pairs = (labels.get("sheet_pairs") or "Pairs")

    def _df_csv(df, index=False):
        buf = io.StringIO()
        df.to_csv(buf, index=index)
        return buf.getvalue()

    def _rows_csv(rows):
        buf = io.StringIO()
        w = csv.writer(buf)
        for r in rows:
            w.writerow(r)
        return buf.getvalue()

    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{name_overview}.csv",
                    _rows_csv([("Metric", "Value"), *_testing_overview_rows(result)]))

        cm = result.get("confusion")
        if cm is not None and not cm.empty:
            zf.writestr(f"{name_confusion}.csv", _df_csv(cm, index=True))

        per_code = result.get("per_code")
        if per_code is not None and not per_code.empty:
            zf.writestr(f"{name_per_code}.csv", _df_csv(per_code, index=False))

        pairs = result.get("pairs") or []
        pairs_df = pd.DataFrame(
            [{"Impuls": p.get("impuls"), "Code A": p.get("code_a"), "Code B": p.get("code_b")}
             for p in pairs],
            columns=["Impuls", "Code A", "Code B"],
        )
        zf.writestr(f"{name_pairs}.csv", _df_csv(pairs_df, index=False))


def export_testing_agreement_json(output_path, result, labels=None):
    """Dump the full result dict (metrics + confusion + per_code + pairs) as JSON."""
    cm = result.get("confusion")
    per_code = result.get("per_code")

    def _num(v):
        try:
            if v != v:  # NaN check
                return None
        except TypeError:
            pass
        return v

    payload = {
        "kappa": _num(result.get("kappa")),
        "kappa_ci_low": _num(result.get("kappa_ci_low")),
        "kappa_ci_high": _num(result.get("kappa_ci_high")),
        "percent_agreement": _num(result.get("percent_agreement")),
        "krippendorff_alpha": _num(result.get("krippendorff_alpha")),
        "gwet_ac1": _num(result.get("gwet_ac1")),
        "brennan_prediger": _num(result.get("brennan_prediger")),
        "n_pairs": result.get("n_pairs"),
        "n_both": result.get("n_both"),
        "n_only_a": result.get("n_only_a"),
        "n_only_b": result.get("n_only_b"),
        "confusion": (cm.to_dict(orient="index") if cm is not None and not cm.empty else {}),
        "per_code": (per_code.to_dict(orient="records") if per_code is not None and not per_code.empty else []),
        "pairs": result.get("pairs") or [],
    }
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2, default=str)


def _df_to_html_table(df, index=False):
    if df is None or df.empty:
        return "<p><em>—</em></p>"
    return df.to_html(index=index, border=0, classes="tt-table", escape=True)


def export_testing_agreement_html(output_path, result, labels=None):
    """Standalone printable HTML report with the four sections."""
    labels = labels or {}
    h_overview = _html_escape(labels.get("sheet_overview") or "Overview")
    h_confusion = _html_escape(labels.get("sheet_confusion") or "Confusion")
    h_per_code = _html_escape(labels.get("sheet_per_code") or "Per-Code")
    h_pairs = _html_escape(labels.get("sheet_pairs") or "Pairs")
    title = _html_escape(labels.get("title") or "Intercoder Agreement")

    overview_rows = "".join(
        f"<tr><th>{_html_escape(k)}</th><td>{_html_escape(v)}</td></tr>"
        for k, v in _testing_overview_rows(result)
    )

    cm = result.get("confusion")
    per_code = result.get("per_code")
    pairs = result.get("pairs") or []
    pairs_df = pd.DataFrame(
        [{"Impuls": p.get("impuls"), "Code A": p.get("code_a"), "Code B": p.get("code_b")}
         for p in pairs],
        columns=["Impuls", "Code A", "Code B"],
    )

    html = f"""<!doctype html>
<html lang="en"><head><meta charset="utf-8">
<title>{title}</title>
<style>
 body {{ font-family: -apple-system, Segoe UI, Roboto, sans-serif; margin: 2rem; color: #222; }}
 h1 {{ font-size: 1.4rem; }} h2 {{ font-size: 1.1rem; margin-top: 1.6rem; }}
 table.tt-table, table.tt-overview {{ border-collapse: collapse; margin: 0.5rem 0; }}
 table.tt-table th, table.tt-table td,
 table.tt-overview th, table.tt-overview td {{ border: 1px solid #bbb; padding: 4px 8px; font-size: 0.92rem; vertical-align: top; }}
 table.tt-overview th {{ text-align: left; background: #f2f2f2; }}
 table.tt-table th {{ background: #f2f2f2; }}
 @media print {{ body {{ margin: 1cm; }} }}
</style></head><body>
<h1>{title}</h1>
<h2>{h_overview}</h2>
<table class="tt-overview"><tbody>{overview_rows}</tbody></table>
<h2>{h_confusion}</h2>
{_df_to_html_table(cm, index=True)}
<h2>{h_per_code}</h2>
{_df_to_html_table(per_code, index=False)}
<h2>{h_pairs}</h2>
{_df_to_html_table(pairs_df, index=False)}
</body></html>
"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)


def _add_docx_table_from_df(doc, df, index=False):
    if df is None or df.empty:
        doc.add_paragraph("—")
        return
    cols = list(df.columns)
    header = (["" ] + cols) if index else cols
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid"
    hdr_cells = table.rows[0].cells
    for i, name in enumerate(header):
        hdr_cells[i].text = str(name)
    for idx, row in df.iterrows():
        cells = table.add_row().cells
        offset = 0
        if index:
            cells[0].text = str(idx)
            offset = 1
        for j, col in enumerate(cols):
            cells[j + offset].text = "" if pd.isna(row[col]) else str(row[col])


def export_testing_agreement_docx(output_path, result, labels=None):
    """Write a .docx with overview, confusion matrix, per-code metrics, and pairs."""
    try:
        from docx import Document as _Doc
    except ImportError as e:
        raise RuntimeError("docx_unavailable") from e

    labels = labels or {}
    title = labels.get("title") or "Intercoder Agreement"
    h_overview = labels.get("sheet_overview") or "Overview"
    h_confusion = labels.get("sheet_confusion") or "Confusion"
    h_per_code = labels.get("sheet_per_code") or "Per-Code"
    h_pairs = labels.get("sheet_pairs") or "Pairs"

    doc = _Doc()
    doc.add_heading(title, level=1)

    doc.add_heading(h_overview, level=2)
    overview_df = pd.DataFrame(_testing_overview_rows(result), columns=["Metric", "Value"])
    _add_docx_table_from_df(doc, overview_df, index=False)

    doc.add_heading(h_confusion, level=2)
    _add_docx_table_from_df(doc, result.get("confusion"), index=True)

    doc.add_heading(h_per_code, level=2)
    _add_docx_table_from_df(doc, result.get("per_code"), index=False)

    doc.add_heading(h_pairs, level=2)
    pairs = result.get("pairs") or []
    pairs_df = pd.DataFrame(
        [{"Impuls": p.get("impuls"), "Code A": p.get("code_a"), "Code B": p.get("code_b")}
         for p in pairs],
        columns=["Impuls", "Code A", "Code B"],
    )
    _add_docx_table_from_df(doc, pairs_df, index=False)

    doc.save(output_path)


def export_testing_agreement_pdf(output_path, result, labels=None):
    """Write DOCX to a temp file, then convert to PDF via docx2pdf."""
    if sys.platform.startswith("linux"):
        raise RuntimeError("pdf_unavailable_linux")
    tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_docx.close()
    try:
        export_testing_agreement_docx(tmp_docx.name, result, labels=labels)
        try:
            from docx2pdf import convert
        except ImportError as e:
            raise RuntimeError("pdf_unavailable") from e
        try:
            convert(tmp_docx.name, output_path)
        except Exception as e:
            raise RuntimeError("pdf_unavailable") from e
    finally:
        try:
            os.unlink(tmp_docx.name)
        except OSError:
            pass


def export_testing_agreement_any(output_path, result, fmt, labels=None):
    """Dispatch to the right exporter based on `fmt`."""
    fmt = (fmt or "xlsx").lower()
    if fmt == "xlsx":
        labels = labels or {}
        return export_testing_agreement(
            output_path, result,
            sheet_overview=labels.get("sheet_overview", "Overview"),
            sheet_confusion=labels.get("sheet_confusion", "Confusion"),
            sheet_per_code=labels.get("sheet_per_code", "Per-Code"),
            sheet_pairs=labels.get("sheet_pairs", "Pairs"),
        )
    if fmt == "csv":
        return export_testing_agreement_csv_zip(output_path, result, labels=labels)
    if fmt == "json":
        return export_testing_agreement_json(output_path, result, labels=labels)
    if fmt == "html":
        return export_testing_agreement_html(output_path, result, labels=labels)
    if fmt == "docx":
        return export_testing_agreement_docx(output_path, result, labels=labels)
    if fmt == "pdf":
        return export_testing_agreement_pdf(output_path, result, labels=labels)
    raise ValueError(f"unsupported_format: {fmt}")
