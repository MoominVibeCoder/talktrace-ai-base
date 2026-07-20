"""Rück-Import einer Codierungs-Tabelle aus einem exportierten Report.

Typischer Weg: Report exportieren → in Word/Excel die Codierungen von Hand
korrigieren → hier wieder einlesen, damit das Feedback auf der geprüften
Fassung beruht statt auf der Roh-Ausgabe des Modells.

Zwei Entwurfsentscheidungen, die man kennen sollte:

* **Tabellen werden über ihre Spaltenköpfe gefunden, nie über die Position.**
  Der DOCX-Report enthält inzwischen mehrere Tabellen (Verteilung nach
  Sprechergruppe, dann die Turns); ein `tables[0]` wäre schon einmal still
  auf die falsche gelaufen. Erkannt wird sprachübergreifend: ein deutscher
  Report lässt sich in einer englischen App einlesen und umgekehrt.
* **Konfidenzwerte werden verworfen.** Nach einer Handkorrektur gehört die
  Zahl des Modells nicht mehr zum Code ("EN (92 %)" → "H" heißt nicht, dass
  H zu 92 % passt). Die Codes zählen, die Konfidenz nicht — das Feedback
  braucht sie ohnehin nicht.

PDF wird bewusst NICHT unterstützt: es ist kein Bearbeitungsformat, und aus
einem PDF ließe sich die Zellstruktur nur raten.
"""
import csv
import io
import re
import zipfile
from html.parser import HTMLParser

import pandas as pd

from ..localization.translation import TRANSLATIONS
from .qualitative import strip_confidence

SUPPORTED_SUFFIXES = (".docx", ".xlsx", ".csv", ".zip", ".html", ".htm")


class ReportImportError(Exception):
    """Der Report ließ sich nicht als Codierungs-Tabelle lesen."""


# Reports einer früheren Version tragen hinter dem Code noch Konfidenz-
# Zeichen ("EN (92 %) ●●●"). Die gibt es in freier Wildbahn, also müssen
# sie beim Import wieder abfallen — sonst landet "EN ●●●" als eigener Code
# im Profil.
_CODE_MARK_RE = re.compile(r"[●○••●○]+")


def _clean_code(value) -> str:
    """Reiner Code aus einer Report-Zelle: ohne Konfidenz, ohne Marker."""
    return _CODE_MARK_RE.sub("", strip_confidence(value)).strip()


def _norm(text) -> str:
    """Vergleichsform eines Spaltenkopfs: klein, ohne Mehrfach-Leerzeichen."""
    return re.sub(r"\s+", " ", str(text)).strip().lower()


def _candidates(*keys) -> set:
    """Alle lokalisierten Schreibweisen eines report-Keys, beide Sprachen."""
    out = set()
    for lang in ("de", "en"):
        for key in keys:
            value = TRANSLATIONS.get(lang, {}).get("report", {}).get(key)
            if value:
                out.add(_norm(value))
    return out


def _is_code_header(header) -> bool:
    """Code-Spalte? Deckt "Shortcode", "Shortcode 1", "Code" in beiden
    Sprachen ab — die Nummerierung hängt am Multi-Coding-Modus."""
    h = _norm(header)
    bases = _candidates("shortcode", "code")
    return any(h == b or h.startswith(f"{b} ") for b in bases)


def _match_headers(headers):
    """Ordnet Spaltenköpfe den gebrauchten Rollen zu.

    Gibt (speaker_idx, statement_idx, [code_idx, ...]) zurück oder None,
    wenn die Tabelle keine Codierungs-Tabelle ist.
    """
    speaker_names = _candidates("speaker")
    statement_names = _candidates("teacher_statement")
    speaker_idx = statement_idx = None
    code_idx = []
    for i, h in enumerate(headers):
        hn = _norm(h)
        if speaker_idx is None and hn in speaker_names:
            speaker_idx = i
        elif statement_idx is None and hn in statement_names:
            statement_idx = i
        elif _is_code_header(h):
            code_idx.append(i)
    if statement_idx is None or not code_idx:
        return None
    return speaker_idx, statement_idx, code_idx


def _rows_to_frame(headers, rows):
    """Zeilen einer erkannten Tabelle in die Analyse-Form bringen.

    Ergebnis-Spalten: Sprecher / Impuls / Shortcode — genau das, was
    teacher_code_profile und der Metriken-Pfad erwarten. Es entsteht EINE
    Zeile je Turn mit dem primären Code; Nebencodes aus weiteren Spalten
    werden verworfen, wie in jeder anderen Häufigkeits-Auswertung auch.
    Uncodierte Turns bleiben erhalten (leerer Shortcode) — sie zählen in
    die Redeanteile.
    """
    matched = _match_headers(headers)
    if matched is None:
        return None
    speaker_idx, statement_idx, code_idx = matched
    primary = code_idx[0]
    records = []
    for row in rows:
        if len(row) <= max(statement_idx, primary):
            continue  # unvollständige Zeile (z. B. verbundene Zellen)
        statement = str(row[statement_idx]).strip()
        if not statement:
            continue
        speaker = (str(row[speaker_idx]).strip()
                   if speaker_idx is not None and len(row) > speaker_idx else "")
        code = _clean_code(row[primary])
        records.append({"Sprecher": speaker, "Impuls": statement,
                        "Shortcode": code})
    if not records:
        return None
    return pd.DataFrame(records, columns=["Sprecher", "Impuls", "Shortcode"])


# --- Format-Leser ---------------------------------------------------------
# Jeder gibt den Frame zurück oder None, wenn er nichts Passendes findet.

def _from_docx(path):
    from docx import Document

    doc = Document(path)
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [c.text for c in table.rows[0].cells]
        frame = _rows_to_frame(
            headers, [[c.text for c in r.cells] for r in table.rows[1:]]
        )
        if frame is not None:
            return frame
    return None


def _from_xlsx(path):
    try:
        sheets = pd.read_excel(path, sheet_name=None, dtype=str)
    except ImportError as e:                      # openpyxl fehlt
        raise ReportImportError("xlsx_unavailable") from e
    for frame in sheets.values():
        got = _frame_from_dataframe(frame)
        if got is not None:
            return got
    return None


def _from_csv_zip(path):
    """CSV-Bündel (ZIP) oder eine einzelne CSV-Datei."""
    if zipfile.is_zipfile(path):
        with zipfile.ZipFile(path) as zf:
            for name in zf.namelist():
                if not name.lower().endswith(".csv"):
                    continue
                with zf.open(name) as fh:
                    got = _frame_from_csv_bytes(fh.read())
                if got is not None:
                    return got
        return None
    with open(path, "rb") as fh:
        return _frame_from_csv_bytes(fh.read())


def _frame_from_csv_bytes(raw):
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        return None
    rows = list(csv.reader(io.StringIO(text)))
    if len(rows) < 2:
        return None
    return _rows_to_frame(rows[0], rows[1:])


def _frame_from_dataframe(frame):
    """Bereits geparste Tabelle (XLSX) auf dieselbe Erkennung schicken."""
    if frame is None or frame.empty:
        return None
    headers = [str(c) for c in frame.columns]
    rows = frame.fillna("").astype(str).values.tolist()
    return _rows_to_frame(headers, rows)


class _TableCollector(HTMLParser):
    """Sammelt alle <table>-Inhalte als Listen von Zeilen."""

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.tables = []
        self._rows = None
        self._row = None
        self._cell = None

    def handle_starttag(self, tag, attrs):
        if tag == "table":
            self._rows = []
        elif tag == "tr" and self._rows is not None:
            self._row = []
        elif tag in ("td", "th") and self._row is not None:
            self._cell = []

    def handle_endtag(self, tag):
        if tag == "table" and self._rows is not None:
            self.tables.append(self._rows)
            self._rows = None
        elif tag == "tr" and self._row is not None:
            self._rows.append(self._row)
            self._row = None
        elif tag in ("td", "th") and self._cell is not None:
            self._row.append("".join(self._cell).strip())
            self._cell = None

    def handle_data(self, data):
        if self._cell is not None:
            self._cell.append(data)


def _from_html(path):
    with open(path, "r", encoding="utf-8", errors="replace") as fh:
        parser = _TableCollector()
        parser.feed(fh.read())
    for rows in parser.tables:
        if len(rows) < 2:
            continue
        frame = _rows_to_frame(rows[0], rows[1:])
        if frame is not None:
            return frame
    return None


_READERS = {
    ".docx": _from_docx,
    ".xlsx": _from_xlsx,
    ".csv": _from_csv_zip,
    ".zip": _from_csv_zip,
    ".html": _from_html,
    ".htm": _from_html,
}


def parse_report(path, filename=None):
    """Codierungs-Tabelle aus einem exportierten Report lesen.

    `filename` liefert die Endung, wenn `path` eine Temp-Datei ohne Suffix
    ist (Shiny legt Uploads so ab). Wirft ReportImportError mit einem
    Localization-Key als Nachricht, damit der Handler ihn direkt übersetzen
    kann.
    """
    suffix = _suffix_of(filename or path)
    reader = _READERS.get(suffix)
    if reader is None:
        raise ReportImportError("import_unsupported_format")
    try:
        frame = reader(path)
    except ReportImportError:
        raise
    except Exception as e:                        # defektes / fremdes Dokument
        raise ReportImportError("import_unreadable") from e
    if frame is None or frame.empty:
        raise ReportImportError("import_no_table")
    return frame


def _suffix_of(name) -> str:
    match = re.search(r"(\.[A-Za-z0-9]+)$", str(name or ""))
    return match.group(1).lower() if match else ""


_STUDENT_LABEL_RE = re.compile(r"^s\d{1,3}$", re.IGNORECASE)


def detect_teacher_label(frame, preferred=None):
    """Lehrkraft-Label aus der Sprecher-Spalte ableiten.

    Ohne das rechnen die Kennzahlen still falsch: heißt die Lehrkraft im
    Report "L", die App steht aber auf "LEHRER", fällt in dialog_stats jeder
    Lehrkraft-Turn KOMPLETT aus der Auswertung — der Sprecher-Regex kennt
    nur <lehrkraft>|S<n>, und "L" ist dann keines von beidem. Der Report
    sähe aus, als hätte die Lehrkraft geschwiegen.

    Anders als die konservative Transkript-Erkennung darf hier auch ein
    einzelner Buchstabe gewinnen: die Sprecher stehen in einer eigenen
    Spalte, also ist "alles, was nicht dem Schüler-Muster S<n> folgt"
    eindeutig. `preferred` (die Einstellung der App) gewinnt, sobald sie im
    Report überhaupt vorkommt. None, wenn sich nichts ableiten lässt.
    """
    if frame is None or frame.empty or "Sprecher" not in frame.columns:
        return None
    speakers = frame["Sprecher"].astype(str).str.strip()
    speakers = speakers[speakers != ""]
    if speakers.empty:
        return None
    if preferred:
        wanted = str(preferred).strip().lower()
        hit = speakers[speakers.str.lower() == wanted]
        if not hit.empty:
            return hit.iloc[0]
    non_students = speakers[~speakers.str.match(_STUDENT_LABEL_RE)]
    if non_students.empty:
        return None
    return non_students.value_counts().idxmax()


def turns_to_transcript(frame, teacher_name=None) -> str:
    """Turns als Transkript-Text ("SPRECHER: Äußerung") zusammensetzen.

    Damit laufen die vorhandenen Statistikfunktionen (dialog_stats &
    Verwandte) unverändert auf einem importierten Report — die Kennzahlen
    entstehen auf demselben Weg wie im normalen Analyse-Lauf, statt hier ein
    zweites Mal nachgebaut zu werden.
    """
    if frame is None or frame.empty:
        return ""
    lines = []
    for _, row in frame.iterrows():
        speaker = str(row.get("Sprecher", "")).strip() or (teacher_name or "")
        utterance = re.sub(r"\s+", " ", str(row.get("Impuls", ""))).strip()
        if speaker and utterance:
            lines.append(f"{speaker}: {utterance}")
    return "\n".join(lines)


def coded_turn_count(frame) -> int:
    """Anzahl Turns mit Code — für die Rückmeldung nach dem Import."""
    if frame is None or frame.empty or "Shortcode" not in frame.columns:
        return 0
    return int((frame["Shortcode"].astype(str).str.strip() != "").sum())
