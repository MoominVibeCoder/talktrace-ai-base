"""Consent-declaration generator (GDPR Art. 13).

A slim, training-context-specific consent form for the case where a
TalkTrace trainer team works *with* teachers: each teacher consents to
having their *own* spoken words recorded, transcribed locally and analysed.

The text is adapted from the CC0-licensed "Consent-Gen-RDMO" of the TU
Dortmund FDM group (Bernd Zey, Wibke Kleina) — https://github.com/berndzey/Consent-Gen-RDMO.
Because that project is published under CC0 1.0 we may reuse and adapt the
wording freely; this is an aid, not legal advice (see the document footer).

This module is intentionally pure: ``build_consent_html`` takes a plain
data dict plus a resolved strings dict (the current-language "consent"
localization section) and returns HTML. No Shiny, no reactivity here, so
it stays trivially unit-testable.
"""
from __future__ import annotations

import html


# Known LLM providers mapped to a human-readable "Name, Country" used to
# pre-fill the recipient field. Cloud providers outside the EU make the
# transcript a third-country transfer — the consent text must say so.
PROVIDER_LOCATIONS = {
    # Custom endpoint: recipient/location unknown to the app — leave the
    # field empty so the mandatory-field guard flags it for manual entry.
    "custom": "",
    "localmind": "LocalMind, Österreich (EU)",
    "openai": "OpenAI, USA",
    "anthropic": "Anthropic, USA",
    "groq": "Groq, USA",
    "openrouter": "OpenRouter, USA",
    "deepseek": "DeepSeek, China",
    "mistral": "Mistral AI, Frankreich (EU)",
    "ollama": "",  # local
}

# Providers we treat as "local / no transfer" → default the mode to local.
LOCAL_PROVIDERS = {"ollama"}


def default_provider_label(api: str | None) -> str:
    """Display label for the configured provider, or empty for local."""
    if not api:
        return ""
    # Any user-registered custom endpoint (slug ``custom:<id>``): recipient and
    # location are unknown to the app — leave empty so the mandatory-field guard
    # flags it for manual entry (same as the former single custom endpoint).
    if api.lower().startswith("custom:"):
        return ""
    return PROVIDER_LOCATIONS.get(api.lower(), api)


def default_llm_mode(api: str | None) -> str:
    """'local' for on-device providers, otherwise 'cloud'."""
    if api and api.lower() in LOCAL_PROVIDERS:
        return "local"
    return "cloud"


# Field keys the form collects and the document consumes. Kept here so the
# handler and the builder agree on one source of truth.
FIELDS = (
    "project_name", "responsible", "dpo", "purpose", "legal_basis",
    "data_categories", "storage", "llm_mode", "llm_provider",
    "revocation", "authority", "participant_name", "place_date",
)

# The eight data-subject-right localization keys, in display order.
RIGHTS_KEYS = (
    "right_withdraw", "right_access", "right_rectify", "right_erase",
    "right_restrict", "right_portability", "right_object", "right_complaint",
)

# Print-friendly stylesheet. Scoped to .ttai-consent-doc so the in-app
# preview doesn't bleed into the rest of the UI; reused verbatim in the
# standalone download (where it also drives the browser's print-to-PDF).
_DOC_CSS = """
.ttai-consent-doc {
  color: #000; background: #fff; font-family: Georgia, "Times New Roman", serif;
  font-size: 11pt; line-height: 1.45; max-width: 48rem; margin: 0 auto;
  padding: 1.5rem 1.75rem; box-sizing: border-box;
}
/* !important: the app's dark-theme heading rules tie on specificity and
   would otherwise win the cascade, washing the headings out to near-white. */
.ttai-consent-doc, .ttai-consent-doc h1, .ttai-consent-doc h2,
.ttai-consent-doc h3, .ttai-consent-doc h4, .ttai-consent-doc p,
.ttai-consent-doc b, .ttai-consent-doc li, .ttai-consent-doc td,
.ttai-consent-doc th { color: #000 !important; }
.ttai-consent-doc h1 { font-size: 17pt; margin: 0 0 0.4rem; }
.ttai-consent-doc h2 { font-size: 13pt; margin: 1.4rem 0 0.5rem;
  border-bottom: 1px solid #555; padding-bottom: 0.15rem; }
.ttai-consent-doc h3 { font-size: 11.5pt; margin: 1rem 0 0.3rem; }
/* Justify the longer prose blocks; short labels/checkboxes stay left. */
.ttai-consent-doc p { margin: 0.5rem 0; text-align: justify; }
.ttai-consent-doc .ttai-checkbox, .ttai-consent-doc .ttai-sigline { text-align: left; }
.ttai-consent-doc ul { margin: 0.4rem 0; padding-left: 1.4rem; }
.ttai-consent-doc li { margin: 0.25rem 0; }
.ttai-consent-doc .ttai-field { white-space: pre-wrap; }
.ttai-consent-doc .ttai-checkbox { font-size: 13pt; letter-spacing: 0.3rem; margin: 0.3rem 0; }
.ttai-consent-doc .ttai-siglines { margin-top: 2rem; }
.ttai-consent-doc .ttai-sigline { border-top: 1px solid #333;
  width: 22rem; margin-top: 2.2rem; padding-top: 0.2rem; font-size: 9.5pt; color: #444; }
.ttai-consent-doc .ttai-footer { margin-top: 2rem; padding-top: 0.5rem;
  border-top: 1px solid #ccc; font-size: 8.5pt; color: #666; font-style: italic; }
.ttai-consent-doc .ttai-missing { color: #b00; font-weight: 700; }
@media print {
  body { background: #fff; }
  .ttai-consent-doc { max-width: none; margin: 0; padding: 0; }
}
"""


# Inline !important on headings: the in-app preview renders inside the
# Bootstrap dark theme, whose heading colour sits in a CSS @layer. Layered
# !important can outrank an unlayered stylesheet !important, so only an
# inline !important is guaranteed to keep the headings black in the preview.
_HSTYLE = ' style="color:#000 !important"'


def _esc(value: str | None) -> str:
    return html.escape((value or "").strip())


def _field_or_missing(value: str | None, missing_label: str) -> str:
    """Escaped value, or a red '!!! …' placeholder when the user left a
    required field empty — mirrors the RDMO template's own behaviour so a
    half-filled form is obviously incomplete on paper."""
    v = (value or "").strip()
    if v:
        return f'<span class="ttai-field">{html.escape(v)}</span>'
    return f'<span class="ttai-missing">!!! {html.escape(missing_label)} !!!</span>'


def _opt_block(label: str, value: str | None) -> str:
    """A '<b>label</b>: value' paragraph, omitted entirely if value empty."""
    v = (value or "").strip()
    if not v:
        return ""
    return f"<p><b>{html.escape(label)}</b><br>{html.escape(v)}</p>"


def build_consent_html(data: dict, S: dict) -> str:
    """Render the inner consent-document HTML (no <html> wrapper).

    ``data``  — form values keyed by :data:`FIELDS`.
    ``S``     — the current-language "consent" localization section.
    """
    g = data.get
    is_cloud = (g("llm_mode") or "cloud") == "cloud"
    provider = (g("llm_provider") or "").strip()
    project = (g("project_name") or "").strip()
    yes, no = S["doc_yes"], S["doc_no"]

    parts: list[str] = [f'<div class="ttai-consent-doc">']

    # --- Title --------------------------------------------------------
    title = S["doc_title"]
    if project:
        title = f"{title} – {project}"
    parts.append(f"<h1{_HSTYLE}>{html.escape(title)}</h1>")

    # --- Part 1: study information ------------------------------------
    parts.append(f"<h2{_HSTYLE}>{html.escape(S['doc_part1_heading'])}</h2>")
    intro = S["doc_intro"].format(project=html.escape(project) if project else S["doc_project_generic"])
    parts.append(f"<p>{intro}</p>")
    parts.append(f"<p>{html.escape(S['doc_pipeline'])}</p>")

    # --- Part 2: Art. 13 information ----------------------------------
    parts.append(f"<h2{_HSTYLE}>{html.escape(S['doc_part2_heading'])}</h2>")
    parts.append(f"<p>{html.escape(S['doc_art13_intro'])}</p>")

    parts.append(
        f"<p><b>{html.escape(S['doc_responsible_label'])}</b><br>"
        f"{_field_or_missing(g('responsible'), S['missing_responsible'])}</p>"
    )
    parts.append(_opt_block(S["doc_dpo_label"], g("dpo")))

    parts.append(
        f"<p><b>{html.escape(S['doc_purpose_label'])}</b><br>"
        f"{_field_or_missing(g('purpose'), S['missing_purpose'])}</p>"
    )
    parts.append(
        f"<p><b>{html.escape(S['doc_legal_basis_label'])}</b><br>"
        f"{_field_or_missing(g('legal_basis'), S['missing_legal_basis'])}</p>"
    )
    parts.append(
        f"<p><b>{html.escape(S['doc_data_categories_label'])}</b><br>"
        f"{_field_or_missing(g('data_categories'), S['missing_data_categories'])}</p>"
    )

    # Recipients / processing — the privacy crux: local transcription vs.
    # cloud-LLM transfer.
    parts.append(f"<p><b>{html.escape(S['doc_recipients_label'])}</b></p>")
    parts.append(f"<p>{html.escape(S['doc_transcription_note'])}</p>")
    if is_cloud:
        prov = html.escape(provider) if provider else html.escape(S["doc_provider_generic"])
        parts.append(f"<p>{S['doc_recipients_cloud'].format(provider=prov)}</p>")
    else:
        parts.append(f"<p>{html.escape(S['doc_recipients_local'])}</p>")

    parts.append(
        f"<p><b>{html.escape(S['doc_storage_label'])}</b><br>"
        f"{S['doc_storage_sentence'].format(duration=_field_or_missing(g('storage'), S['missing_storage']))}</p>"
    )

    # Rights
    parts.append(f"<p><b>{html.escape(S['doc_rights_label'])}</b></p><ul>")
    for key in RIGHTS_KEYS:
        # values are 'Label: text'; bold the label up to the first colon.
        raw = S[key]
        if ":" in raw:
            lbl, rest = raw.split(":", 1)
            parts.append(f"<li><b>{html.escape(lbl)}:</b>{html.escape(rest)}</li>")
        else:
            parts.append(f"<li>{html.escape(raw)}</li>")
    parts.append("</ul>")
    parts.append(
        f"<p>{S['doc_rights_contact'].format(responsible=_field_or_missing(g('responsible'), S['missing_responsible']))}</p>"
    )
    parts.append(
        f"<p>{S['doc_revocation_sentence'].format(revocation=_field_or_missing(g('revocation'), S['missing_revocation']))}</p>"
    )
    parts.append(_opt_block(S["doc_authority_label"], g("authority")))

    # --- Part 3: consent declaration ----------------------------------
    parts.append(f"<h2{_HSTYLE}>{html.escape(S['doc_part3_heading'])}</h2>")
    parts.append(f"<p>{html.escape(S['doc_voluntary'])}</p>")

    parts.append(f"<p>{S['doc_consent_main'].format(purpose=html.escape((g('purpose') or '').strip()) or S['doc_purpose_generic'])}</p>")
    parts.append(f'<p class="ttai-checkbox">&#x2610; {html.escape(yes)} &nbsp;&nbsp; &#x2610; {html.escape(no)}</p>')

    if is_cloud:
        prov = html.escape(provider) if provider else html.escape(S["doc_provider_generic"])
        parts.append(f"<p>{S['doc_consent_cloud'].format(provider=prov)}</p>")
        parts.append(f'<p class="ttai-checkbox">&#x2610; {html.escape(yes)} &nbsp;&nbsp; &#x2610; {html.escape(no)}</p>')

    if project:
        parts.append(f"<p>{S['doc_consent_participate'].format(project=html.escape(project))}</p>")

    # Signature block
    parts.append('<div class="ttai-siglines">')
    name = (g("participant_name") or "").strip()
    place = (g("place_date") or "").strip()
    parts.append(
        f'<div class="ttai-sigline">{html.escape(name) if name else ""}<br>{html.escape(S["doc_name_line"])}</div>'
    )
    parts.append(
        f'<div class="ttai-sigline">{html.escape(place) if place else ""}<br>{html.escape(S["doc_signature_line"])}</div>'
    )
    parts.append("</div>")

    # Footer: CC0 credit + disclaimer
    parts.append(f'<div class="ttai-footer">{html.escape(S["doc_footer"])}</div>')

    parts.append("</div>")
    return "".join(parts)


def build_consent_preview(data: dict, S: dict) -> str:
    """Inner document HTML plus the scoped <style> — for the in-app preview."""
    return f"<style>{_DOC_CSS}</style>{build_consent_html(data, S)}"


def write_consent_docx(path: str, data: dict, S: dict) -> None:
    """Render the consent declaration into a Word .docx at ``path``.

    Mirrors :func:`build_consent_html`: black headings, justified prose,
    bold field labels, bullet rights, '☐ Ja ☐ Nein' checkboxes, signature
    lines, and red '!!! … !!!' markers for missing required fields. Kept in
    sync with the HTML builder by hand — the content here is intentionally
    the same wording, just emitted through python-docx.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY
    BLACK = RGBColor(0x00, 0x00, 0x00)
    RED = RGBColor(0xB0, 0x00, 0x00)

    g = data.get
    is_cloud = (g("llm_mode") or "cloud") == "cloud"
    provider = (g("llm_provider") or "").strip()
    project = (g("project_name") or "").strip()
    yes, no = S["doc_yes"], S["doc_no"]

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def heading(text, size):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = BLACK
        return p

    def prose(text):
        p = doc.add_paragraph(text)
        p.alignment = JUSTIFY
        return p

    def missing_run(p, label):
        r = p.add_run(f"!!! {label} !!!")
        r.bold = True
        r.font.color.rgb = RED

    def label_value(label, value, missing):
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p2 = doc.add_paragraph()
        p2.alignment = JUSTIFY
        v = (value or "").strip()
        if v:
            p2.add_run(v)
        else:
            missing_run(p2, missing)
        return p2

    def opt_block(label, value):
        v = (value or "").strip()
        if not v:
            return
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p2 = doc.add_paragraph(v)
        p2.alignment = JUSTIFY

    def templated(template, key, value, missing):
        """A justified paragraph 'before {key} after' where the slot is the
        value, or a red missing-marker when empty."""
        before, _, after = template.partition("{" + key + "}")
        p = doc.add_paragraph()
        p.alignment = JUSTIFY
        if before:
            p.add_run(before)
        v = (value or "").strip()
        if v:
            p.add_run(v)
        else:
            missing_run(p, missing)
        if after:
            p.add_run(after)
        return p

    def checkbox():
        doc.add_paragraph(f"☐ {yes}     ☐ {no}")

    # --- Title --------------------------------------------------------
    title = S["doc_title"]
    if project:
        title = f"{title} – {project}"
    heading(title, 16)

    # --- Part 1 -------------------------------------------------------
    heading(S["doc_part1_heading"], 13)
    prose(S["doc_intro"].format(project=project or S["doc_project_generic"]))
    prose(S["doc_pipeline"])

    # --- Part 2 -------------------------------------------------------
    heading(S["doc_part2_heading"], 13)
    prose(S["doc_art13_intro"])
    label_value(S["doc_responsible_label"], g("responsible"), S["missing_responsible"])
    opt_block(S["doc_dpo_label"], g("dpo"))
    label_value(S["doc_purpose_label"], g("purpose"), S["missing_purpose"])
    label_value(S["doc_legal_basis_label"], g("legal_basis"), S["missing_legal_basis"])
    label_value(S["doc_data_categories_label"], g("data_categories"), S["missing_data_categories"])

    doc.add_paragraph().add_run(S["doc_recipients_label"]).bold = True
    prose(S["doc_transcription_note"])
    if is_cloud:
        prose(S["doc_recipients_cloud"].format(provider=provider or S["doc_provider_generic"]))
    else:
        prose(S["doc_recipients_local"])

    doc.add_paragraph().add_run(S["doc_storage_label"]).bold = True
    templated(S["doc_storage_sentence"], "duration", g("storage"), S["missing_storage"])

    doc.add_paragraph().add_run(S["doc_rights_label"]).bold = True
    for key in RIGHTS_KEYS:
        raw = S[key]
        p = doc.add_paragraph(style="List Bullet")
        if ":" in raw:
            lbl, rest = raw.split(":", 1)
            p.add_run(lbl + ":").bold = True
            p.add_run(rest)
        else:
            p.add_run(raw)
    templated(S["doc_rights_contact"], "responsible", g("responsible"), S["missing_responsible"])
    templated(S["doc_revocation_sentence"], "revocation", g("revocation"), S["missing_revocation"])
    opt_block(S["doc_authority_label"], g("authority"))

    # --- Part 3 -------------------------------------------------------
    heading(S["doc_part3_heading"], 13)
    prose(S["doc_voluntary"])
    prose(S["doc_consent_main"].format(purpose=(g("purpose") or "").strip() or S["doc_purpose_generic"]))
    checkbox()
    if is_cloud:
        prose(S["doc_consent_cloud"].format(provider=provider or S["doc_provider_generic"]))
        checkbox()
    if project:
        prose(S["doc_consent_participate"].format(project=project))

    # Signature block
    doc.add_paragraph()
    for value, caption in (
        ((g("participant_name") or "").strip(), S["doc_name_line"]),
        ((g("place_date") or "").strip(), S["doc_signature_line"]),
    ):
        if value:
            doc.add_paragraph(value)
        doc.add_paragraph("_" * 45)
        cap = doc.add_paragraph()
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)

    # Footer
    foot = doc.add_paragraph()
    fr = foot.add_run(S["doc_footer"])
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.save(path)


def build_consent_standalone(data: dict, S: dict) -> str:
    """A complete, self-contained HTML document for download/print."""
    title = S["doc_title"]
    project = (data.get("project_name") or "").strip()
    if project:
        title = f"{title} – {project}"
    return (
        "<!DOCTYPE html>\n"
        f'<html lang="{S.get("doc_lang", "de")}">\n<head>\n'
        '<meta charset="utf-8">\n'
        '<meta name="viewport" content="width=device-width, initial-scale=1">\n'
        f"<title>{html.escape(title)}</title>\n"
        f"<style>{_DOC_CSS}\nbody{{margin:0;background:#fff;}}</style>\n"
        "</head>\n<body>\n"
        f"{build_consent_html(data, S)}\n"
        "</body>\n</html>\n"
    )
