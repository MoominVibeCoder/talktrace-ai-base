"""Tests für die T-SEDA-Vorlage und die Konfidenz-Pipeline (Multi-Coding).

Deckt die pure Logik ab: Vorlagen-Inhalt (Codes, DE/EN, Auffang-Kategorie
zuletzt), Schema-Erweiterung (Konfidenz nullable + required), Item-
Normalisierung, DataFrame-Aufbau und das Top-3/Cutoff-Postprocessing der
Ergebnistabelle.
"""
import matplotlib
import pandas as pd
import pytest

matplotlib.use("Agg")  # headless: die Plot-Helfer dürfen kein Tk anfassen

from talktrace_ai.examples.tseda import (
    TSEDA_ATTRIBUTION,
    TSEDA_CODEBOOK,
    TSEDA_PRESETS,
)
from talktrace_ai.utils.llm_analysis._schema import (
    build_analysis_schema,
    extract_shortcodes,
)
from talktrace_ai.utils.llm_analysis._json import _extract_items_by_schema
from talktrace_ai.utils.llm_analysis._shared import analysis_items_to_df
from talktrace_ai.utils.llm_analysis._stream_parse import (
    normalize_item,
    parse_jsonl_line,
)
from talktrace_ai.utils.qualitative import (
    MAX_CODES_PER_TURN,
    aggregate_multicoded,
    build_qual_plot,
    build_qual_stats_df,
    code_column_names,
    code_counts_by_group,
    CONFIDENCE_HIGH_MIN,
    CONFIDENCE_LOW_MAX,
    confidence_band,
    confidence_band_of_cell,
    extract_confidence,
    collect_codes,
    primary_code_over_time,
    primary_code_series,
    strip_confidence,
    uncoded_turns,
)


# ---------------------------------------------------------------------------
# T-SEDA-Vorlage
# ---------------------------------------------------------------------------

# Reihenfolge = offizielle Detailfassung (Pack-Folien 33–38); der Crosswalk
# spiegelt sie 1:1 (dt. R=Reflexion ↔ engl. RD; dt. N=Begründen ↔ engl. R).
_EXPECTED_DE = ["I", "EI", "H", "N", "EN", "ZK", "R", "V", "L", "ÄN"]
_EXPECTED_EN = ["B", "IB", "CH", "R", "IR", "CA", "RD", "C", "G", "E"]


def test_tseda_codebook_codes():
    assert extract_shortcodes(TSEDA_CODEBOOK["de"]) == _EXPECTED_DE
    assert extract_shortcodes(TSEDA_CODEBOOK["en"]) == _EXPECTED_EN


def test_tseda_codebook_shape_matches_docx_import():
    # Gleiche Form wie ein docx-Tabellen-Import: list[dict] mit konstanten Keys.
    for entry in TSEDA_CODEBOOK["de"]:
        assert set(entry.keys()) == {"Code", "Bezeichnung", "Beschreibung"}
        assert entry["Beschreibung"].strip()
    for entry in TSEDA_CODEBOOK["en"]:
        assert set(entry.keys()) == {"Code", "Label", "Description"}
        assert entry["Description"].strip()


def test_tseda_catchall_is_last():
    # Position = Fallback-Hierarchie: die Auffang-Kategorie muss zuletzt stehen,
    # damit sie im Single-Coding-Modus nie einen spezifischeren Code verdrängt.
    assert TSEDA_CODEBOOK["de"][-1]["Code"] == "ÄN"
    assert TSEDA_CODEBOOK["en"][-1]["Code"] == "E"


def test_tseda_presets():
    # Seit 2026-07 codiert die Vorlage BEIDE Sprechergruppen: T-SEDA erfasst
    # dialogische Züge aller Beteiligten, und Lehrkraft-Züge sind ohne die
    # umgebenden Schülerturns oft nicht interpretierbar.
    assert TSEDA_PRESETS["llm_switch"] is True
    assert TSEDA_PRESETS["analyse_teacher_switch"] is True
    assert TSEDA_PRESETS["analyse_students_switch"] is True
    assert TSEDA_PRESETS["multi_coding_switch"] is True


def test_context_prompt_keys_exist_in_both_languages():
    # Kontext-Anweisung (Turns im Gesprächsverlauf codieren) und
    # Relevanz-Regel (Nicht-Züge wie Drannehmen/Minimal-Feedback/(unv.)
    # bleiben uncodiert) hängen als Suffixe an System- UND User-Prompt —
    # alle Keys müssen in beiden Sprachen existieren und nicht leer sein.
    from talktrace_ai.localization.translation import TRANSLATIONS
    for lang in ("de", "en"):
        for key in ("prompt_context", "user_prompt_context",
                    "prompt_relevance", "user_prompt_relevance"):
            assert TRANSLATIONS[lang]["sidebar"][key].strip()


def test_invite_reasoning_is_narrowed_to_justification():
    # Testrunde 7: Schüler-Verständnisfragen landeten auf EN, weil das
    # offizielle Schlüsselwort „Kannst du das genauer erklären?" funktional
    # eine KLÄRUNGS-Bitte ist (die das Schema unter EI führt). EN meint das
    # Einfordern einer BEGRÜNDUNG; Rückfragen werden nach ihrer Funktion
    # codiert (I/H/N), Organisationsfragen bleiben uncodiert.
    de = {e["Code"]: e["Beschreibung"] for e in TSEDA_CODEBOOK["de"]}["EN"]
    en = {e["Code"]: e["Description"] for e in TSEDA_CODEBOOK["en"]}["IR"]
    assert "BEGRÜNDUNG" in de and "JUSTIFICATION" in en
    # Das mehrdeutige Klärungs-Keyword ist raus …
    assert "genauer erklären" not in de
    assert "explain that in more detail" not in en
    # … und die Umleitung auf I/H/N bzw. B/CH/R ist benannt.
    for token in ("→ I", "→ H", "→ N"):
        assert token in de, token
    for token in ("→ B", "→ CH", "→ R"):
        assert token in en, token
    assert "uncodiert" in de and "uncoded" in en  # Organisationsfragen


def test_relevance_rule_is_scoped_and_takes_precedence():
    # Zwei Präzisierungen aus Testrunde 6, beide sprachübergreifend:
    # (1) Scoping — die Quittungs-Regel gilt nur für Beiträge, die
    #     AUSSCHLIESSLICH daraus bestehen (sonst verlor das Modell
    #     substanzielle Turns, die bloß mit „Ja," beginnen).
    # (2) Vorrang — kein Eintrag mit Niedrig-Konfidenz statt Weglassen.
    from talktrace_ai.localization.translation import TRANSLATIONS
    for key in ("prompt_relevance", "user_prompt_relevance"):
        de = TRANSLATIONS["de"]["sidebar"][key]
        # Bindestrich-tolerant: "low confidence" / "low-confidence".
        en = TRANSLATIONS["en"]["sidebar"][key].replace("-", " ")
        assert "AUSSCHLIESSLICH" in de, key
        assert "EXCLUSIVELY" in en, key
        assert "niedriger Konfidenz" in de, key
        assert "low confidence" in en, key


def test_multi_coding_prompt_has_calibration_anchors():
    # Kalibrier-Anker gegen uniforme 90/95er-Konfidenzen (Befund Testrunde 5):
    # volle Skala mit benannten Stufen, kein pauschaler Standardwert.
    from talktrace_ai.localization.translation import TRANSLATIONS
    for lang in ("de", "en"):
        for key in ("prompt_multi_coding_on", "user_prompt_multi_coding_on"):
            txt = TRANSLATIONS[lang]["sidebar"][key]
            assert "90" in txt and "60" in txt and "50" in txt


def test_tseda_attribution_names_cambridge():
    assert "Cambridge" in TSEDA_ATTRIBUTION
    assert "CC BY" in TSEDA_ATTRIBUTION


def test_teacher_typical_codes_carry_speaker_note():
    # EN/ZK/L (EN: IR/CA/G) sind im Klassengespräch Lehrkraft-Züge — bei
    # Schüler:innen nur ohne beteiligte Lehrkraft plausibel (Befund
    # Testrunde 5: Modelle vergaben Einlade-/Leitungs-Codes an Schülerturns).
    de_map = {e["Code"]: e["Beschreibung"] for e in TSEDA_CODEBOOK["de"]}
    en_map = {e["Code"]: e["Description"] for e in TSEDA_CODEBOOK["en"]}
    for c in ("EN", "ZK", "L"):
        assert "Klassengespräch" in de_map[c], c
    for c in ("IR", "CA", "G"):
        assert "whole-class" in en_map[c], c
    # Die übrigen Codes bleiben sprecherneutral (offizielle Fassung) — EI/IB
    # bewusst darunter (Testrunde 8): eine Klärungsfrage ist kein
    # lehrkraft-privilegierter Zug.
    for c in ("I", "EI", "H", "N", "R", "V", "ÄN"):
        assert "Klassengespräch" not in de_map[c], c
    for c in ("B", "IB", "CH", "R", "RD", "C", "E"):
        assert "whole-class" not in en_map[c], c


def test_invite_to_build_is_explicitly_speaker_neutral():
    # Testrunde 8: die vier Schüler-Klärungsfragen wanderten von EN nach EI —
    # funktional korrekt. EI trägt die Sprecherneutralität jetzt explizit,
    # damit die Sperre nicht über die Lehrkraft-Notiz zurückkommt.
    de = {e["Code"]: e["Beschreibung"] for e in TSEDA_CODEBOOK["de"]}["EI"]
    en = {e["Code"]: e["Description"] for e in TSEDA_CODEBOOK["en"]}["IB"]
    assert "SPRECHERNEUTRAL" in de
    assert "SPEAKER-NEUTRAL" in en


# ---------------------------------------------------------------------------
# Schema: Konfidenz nullable + required (OpenAI-strict-kompatibel)
# ---------------------------------------------------------------------------

def test_schema_has_nullable_required_confidence():
    schema = build_analysis_schema(TSEDA_CODEBOOK["de"], "LEHRER: Hallo.\nS01: Hi.")
    items = schema["properties"]["analysis"]["items"]
    assert items["properties"]["Konfidenz"]["type"] == ["integer", "null"]
    assert "Konfidenz" in items["required"]
    # Konfidenz als LETZTES Feld definiert (Recovery-Parser-Kontrakt).
    assert list(items["properties"].keys())[-1] == "Konfidenz"


# ---------------------------------------------------------------------------
# Parsing: normalize_item / JSONL / Recovery
# ---------------------------------------------------------------------------

def test_normalize_item_passes_confidence_through():
    item = normalize_item({"#": 1, "Sprecher": "LEHRER", "Shortcode": "EN",
                           "Impuls": "Warum?", "Konfidenz": 85})
    assert item["Konfidenz"] == 85


def test_normalize_item_clamps_and_coerces_confidence():
    assert normalize_item({"Shortcode": "EN", "Impuls": "x", "Konfidenz": 250})["Konfidenz"] == 100
    assert normalize_item({"Shortcode": "EN", "Impuls": "x", "Konfidenz": -3})["Konfidenz"] == 0
    assert normalize_item({"Shortcode": "EN", "Impuls": "x", "Konfidenz": "72"})["Konfidenz"] == 72


def test_normalize_item_without_confidence_keeps_legacy_shape():
    item = normalize_item({"#": 1, "Sprecher": "LEHRER", "Shortcode": "EN", "Impuls": "Warum?"})
    assert "Konfidenz" not in item
    # Explizites null (Single-Coding mit strict-Schema) bleibt ebenfalls draußen.
    item = normalize_item({"Shortcode": "EN", "Impuls": "x", "Konfidenz": None})
    assert "Konfidenz" not in item


def test_parse_jsonl_line_with_confidence():
    line = '{"#": 3, "Sprecher": "LEHRER", "Shortcode": "EN", "Impuls": "Warum?", "Konfidenz": 90}'
    assert parse_jsonl_line(line)["Konfidenz"] == 90


def test_recovery_parser_captures_trailing_confidence():
    text = (
        '{"analysis": ['
        '{"#": 1, "Sprecher": "LEHRER", "Shortcode": "EN", "Impuls": "Warum "denn" so?", "Konfidenz": 77},'
        '{"#": 2, "Sprecher": "S01", "Shortcode": "N", "Impuls": "Weil."}'
        ']}'
    )
    items = _extract_items_by_schema(text)
    assert len(items) == 2
    assert items[0]["Konfidenz"] == 77
    assert "Konfidenz" not in items[1]


# ---------------------------------------------------------------------------
# DataFrame-Aufbau
# ---------------------------------------------------------------------------

def test_items_to_df_without_confidence_keeps_four_columns():
    df = analysis_items_to_df([
        {"#": 1, "Sprecher": "LEHRER", "Shortcode": "EN", "Impuls": "Warum?"},
    ])
    assert list(df.columns) == ["#", "Sprecher", "Shortcode", "Impuls"]


def test_items_to_df_with_confidence_adds_column():
    df = analysis_items_to_df([
        {"#": 1, "Sprecher": "LEHRER", "Shortcode": "EN", "Impuls": "Warum?", "Konfidenz": 80},
        {"#": 1, "Sprecher": "LEHRER", "Shortcode": "L", "Impuls": "Warum?"},
    ])
    assert list(df.columns) == ["#", "Sprecher", "Shortcode", "Impuls", "Konfidenz"]
    assert df["Konfidenz"].tolist()[0] == 80


def test_items_to_df_empty():
    df = analysis_items_to_df([])
    assert list(df.columns) == ["#", "Sprecher", "Shortcode", "Impuls"]
    assert df.empty


# ---------------------------------------------------------------------------
# Postprocessing: Top-3-Spalten, Konfidenz-Sortierung, Anzeige-Format
# ---------------------------------------------------------------------------

def _coded(rows):
    """rows: list of (key, code, priority, konfidenz|None)"""
    return pd.DataFrame(
        [{"__key__": k, "Shortcode": c, "__priority__": p, "Konfidenz": v}
         for k, c, p, v in rows]
    ).sort_values("__priority__", kind="mergesort")


def _wide(out):
    return out[["__code1__", "__code2__"]].values.tolist()


def test_aggregate_caps_top2_no_cutoff():
    # Kein Konfidenz-Filter: auch unsichere Kandidaten erscheinen — aber
    # höchstens zwei (T-SEDA-Regel: 0–2 Codes pro Turn), Konfidenz absteigend.
    coded = _coded([
        ("t1", "EN", 4, 95),
        ("t1", "H", 3, 40),     # unter 50 — bleibt grundsätzlich sichtbar …
        ("t1", "EI", 1, 60),
        ("t1", "ÄN", 10, 25),
    ])
    out = aggregate_multicoded(coded)
    # … fällt hier aber dem Top-2-Cap zum Opfer (Plätze 3+4 gekappt).
    assert _wide(out) == [["EN (95 %)", "EI (60 %)"]]
    assert MAX_CODES_PER_TURN == 2


def test_aggregate_shows_uncertain_candidate_below_50():
    coded = _coded([
        ("t1", "EN", 4, 95),
        ("t1", "H", 3, 35),    # unsicher — erscheint trotzdem mit Konfidenz
    ])
    out = aggregate_multicoded(coded)
    assert _wide(out) == [["EN (95 %)", "H (35 %)"]]


def test_aggregate_orders_by_confidence_not_priority():
    coded = _coded([
        ("t1", "EI", 1, 60),   # höchste Codebuch-Priorität, aber geringere Konfidenz
        ("t1", "L", 9, 90),
    ])
    out = aggregate_multicoded(coded)
    assert _wide(out) == [["L (90 %)", "EI (60 %)"]]


def test_aggregate_dedupes_same_code_keeps_highest_confidence():
    coded = _coded([
        ("t1", "EN", 4, 60),
        ("t1", "EN", 4, 88),
    ])
    out = aggregate_multicoded(coded)
    assert _wide(out) == [["EN (88 %)", ""]]


def test_aggregate_without_confidence_column_is_priority_order():
    coded = pd.DataFrame([
        {"__key__": "t1", "Shortcode": "L", "__priority__": 9},
        {"__key__": "t1", "Shortcode": "EI", "__priority__": 1},
    ]).sort_values("__priority__", kind="mergesort")
    out = aggregate_multicoded(coded)
    # Ohne Konfidenz: Prioritäts-Reihenfolge, kein Suffix.
    assert _wide(out) == [["EI", "L"]]


def test_aggregate_rows_without_confidence_sort_last():
    # Altbestand (NaN-Konfidenz) bleibt erhalten, sortiert ans Ende.
    coded = _coded([
        ("t1", "EN", 4, 80),
        ("t1", "V", 7, None),
    ])
    out = aggregate_multicoded(coded)
    assert _wide(out) == [["EN (80 %)", "V"]]


def test_strip_confidence():
    assert strip_confidence("EN (85 %)") == "EN"
    assert strip_confidence("EN (85 %); L (62 %)") == "EN; L"
    assert strip_confidence("EN") == "EN"


def test_collect_and_primary_from_wide_table():
    cols = code_column_names(_t)
    assert len(cols) == 2
    df = pd.DataFrame({
        "#": [1, 2],
        cols[0]: ["EN (92 %)", ""],
        cols[1]: ["L (40 %)", ""],
    })
    assert sorted(collect_codes(df, _t).tolist()) == ["EN", "L"]
    assert primary_code_series(df, _t).tolist() == ["EN", ""]


def test_collect_and_primary_from_legacy_table():
    sc = _t("report", "shortcode")
    df = pd.DataFrame({"#": [1, 2], sc: ["EN (92 %); L (40 %)", ""]})
    assert sorted(collect_codes(df, _t).tolist()) == ["EN", "L"]
    assert primary_code_series(df, _t).tolist() == ["EN", ""]


def test_qual_plot_counts_primary_code_only():
    # Das Häufigkeits-Diagramm zählt nur Shortcode 1 — Nebenkandidaten
    # (Spalte 2) verzerren die Verteilung nicht.
    cols = code_column_names(_t)
    df = pd.DataFrame({
        "#": [1, 2],
        _t("report", "speaker"): ["LEHRER", "LEHRER"],
        _t("report", "teacher_statement"): ["a", "b"],
        cols[0]: ["EN (92 %)", "H (80 %)"],
        cols[1]: ["L (61 %)", ""],
    })
    ax = build_qual_plot(df, _t)
    labels = sorted(tick.get_text() for tick in ax.get_xticklabels())
    assert labels == ["EN", "H"]  # L (Zweitcode) taucht nicht auf


def test_primary_code_over_time_counts_primary_only():
    # Der Zeitverlauf zählt je Turn nur Shortcode 1 — wie Balken/Chip/
    # Übergangsmatrix. Der Sekundärcode (Spalte 2) darf nirgends auftauchen.
    c1, c2 = code_column_names(_t)
    spk = _t("report", "speaker")
    utt = _t("report", "teacher_statement")
    merged = pd.DataFrame({
        "#": range(1, 7),
        spk: ["LEHRER"] * 6,
        utt: [f"u{i}" for i in range(6)],
        # 3 Abschnitte à 2 Turns (bucket_size = 6 // 3 = 2).
        c1: ["EN (90 %)", "EN (80 %)", "H (70 %)", "", "EN (60 %)", "H (55 %)"],
        c2: ["L (50 %)", "", "", "", "", ""],
    })
    dist = primary_code_over_time(merged, _t, n_segments=3)
    # Sekundärcode L taucht nirgends im Verlauf auf.
    assert "L" not in dist["Shortcode"].tolist()
    # Abschnitt 1: nur EN (2×), Anteil 1.0.
    seg1 = dist[dist["Abschnitt"] == "1"]
    assert seg1["Shortcode"].tolist() == ["EN"]
    assert seg1["Anteil"].tolist() == [1.0]
    # Abschnitt 2: nur H (der zweite Turn ist uncodiert und zählt nicht).
    seg2 = dist[dist["Abschnitt"] == "2"].set_index("Shortcode")["Anzahl"].to_dict()
    assert seg2 == {"H": 1}
    # Abschnitt 3: EN und H je einmal.
    seg3 = dist[dist["Abschnitt"] == "3"].set_index("Shortcode")["Anzahl"].to_dict()
    assert seg3 == {"EN": 1, "H": 1}


def test_primary_code_over_time_empty():
    assert primary_code_over_time(pd.DataFrame(), _t).empty


# ---------------------------------------------------------------------------
# Zweite Prüfrunde: Auswahl der uncodierten, codierbaren Turns
# ---------------------------------------------------------------------------

_TRANSCRIPT_2P = (
    "LEHRER: Warum sollte das so sein?\n"
    "S1: Weil alle betroffen sind.\n"
    "LEHRER: S2\n"
    "S2: Ich stimme zu.\n"
    "LEHRER: Fassen wir zusammen.\n"
)


def test_uncoded_turns_teacher_only():
    items = [{"Sprecher": "LEHRER", "Shortcode": "EN",
              "Impuls": "Warum sollte das so sein?"}]
    pending = uncoded_turns(_TRANSCRIPT_2P, "LEHRER", items,
                            teacher_on=True, students_on=False)
    # Nur die uncodierten LEHRER-Turns — keine Schülerturns, der codierte
    # Turn fehlt ebenfalls.
    assert pending == [("LEHRER", "S2"), ("LEHRER", "Fassen wir zusammen.")]


def test_uncoded_turns_matches_tolerantly_and_via_alias():
    # LLM-Item mit generischem Sprecher-Label + Mini-Textabweichung
    # (fehlendes Satzzeichen) muss dem Transkript-Turn zugeordnet werden.
    items = [{"Sprecher": "Lehrperson", "Shortcode": "ZK",
              "Impuls": "Fassen wir zusammen"}]
    pending = uncoded_turns(_TRANSCRIPT_2P, "LEHRER", items,
                            teacher_on=True, students_on=False)
    assert ("LEHRER", "Fassen wir zusammen.") not in pending
    assert ("LEHRER", "Warum sollte das so sein?") in pending


def test_uncoded_turns_students_included_when_enabled():
    pending = uncoded_turns(_TRANSCRIPT_2P, "LEHRER", [],
                            teacher_on=True, students_on=True)
    assert ("S1", "Weil alle betroffen sind.") in pending
    assert ("S2", "Ich stimme zu.") in pending


def test_uncoded_turns_empty_transcript():
    assert uncoded_turns("", "LEHRER", []) == []


# ---------------------------------------------------------------------------
# End-to-End über build_qual_stats_df (pure Spiegelung der Ergebnistabelle)
# ---------------------------------------------------------------------------

def _t(section, key):
    from talktrace_ai.localization.translation import TRANSLATIONS
    return TRANSLATIONS["de"][section][key]


def test_build_qual_stats_df_multicoding_wide_columns():
    # Einstelliges Schülerlabel (S1) mit Absicht: der Parser-Fix (S\d{1,3})
    # muss den Turn in die Tabelle bringen.
    transcript = "LEHRER: Warum sollte das so sein?\nS1: Weil alle betroffen sind."
    analysis_df = analysis_items_to_df([
        {"#": 1, "Sprecher": "LEHRER", "Shortcode": "EN",
         "Impuls": "Warum sollte das so sein?", "Konfidenz": 92},
        {"#": 1, "Sprecher": "LEHRER", "Shortcode": "L",
         "Impuls": "Warum sollte das so sein?", "Konfidenz": 61},
        {"#": 1, "Sprecher": "LEHRER", "Shortcode": "ÄN",
         "Impuls": "Warum sollte das so sein?", "Konfidenz": 30},
    ])
    merged = build_qual_stats_df(
        analysis_df, transcript, "LEHRER", TSEDA_CODEBOOK["de"],
        multi_coding=True, t=_t,
    )
    c1, c2 = code_column_names(_t)
    # Alle Turns des Gesprächs erscheinen — auch der (uncodierte) S1-Turn.
    assert len(merged) == 2
    # Top-2-Kandidaten sichtbar (Konfidenz absteigend); Platz 3 (ÄN, 30 %)
    # fällt dem Cap zum Opfer.
    assert merged[c1].tolist() == ["EN (92 %)", ""]
    assert merged[c2].tolist() == ["L (61 %)", ""]


def test_build_qual_stats_df_single_coding_ignores_confidence():
    transcript = "LEHRER: Warum sollte das so sein?"
    analysis_df = analysis_items_to_df([
        {"#": 1, "Sprecher": "LEHRER", "Shortcode": "L",
         "Impuls": "Warum sollte das so sein?", "Konfidenz": 90},
        {"#": 1, "Sprecher": "LEHRER", "Shortcode": "EN",
         "Impuls": "Warum sollte das so sein?", "Konfidenz": 70},
    ])
    merged = build_qual_stats_df(
        analysis_df, transcript, "LEHRER", TSEDA_CODEBOOK["de"],
        multi_coding=False, t=_t,
    )
    sc_col = _t("report", "shortcode")
    # Single-Coding bleibt beim Hierarchie-Kontrakt: EN steht im Codebuch
    # vor L und gewinnt unabhängig von der Konfidenz — kein Suffix, und
    # die klassische Einzel-Spalte (keine Code-1..3-Spalten).
    assert merged[sc_col].tolist() == ["EN"]
    assert not any(c in merged.columns for c in code_column_names(_t))


# ---------------------------------------------------------------------------
# Code-Verteilung nach Sprechergruppe (Balkenplot + Report-Tabelle)
# ---------------------------------------------------------------------------

def _grouped_df(speakers, primary, secondary=None):
    """Ergebnistabelle in Multi-Coding-Form mit frei wählbaren Sprechern."""
    c1, c2 = code_column_names(_t)
    return pd.DataFrame({
        "#": range(1, len(speakers) + 1),
        _t("report", "speaker"): speakers,
        _t("report", "teacher_statement"): [f"u{i}" for i in range(len(speakers))],
        c1: primary,
        c2: secondary if secondary is not None else [""] * len(speakers),
    })


def test_code_counts_by_group_splits_teacher_and_students():
    df = _grouped_df(
        ["LEHRER", "S1", "LEHRER", "S2"],
        ["EN (90 %)", "EN (80 %)", "L (70 %)", "I (60 %)"],
    )
    counts = code_counts_by_group(df, _t, "LEHRER")
    teacher, pupils = _t("report", "teacher"), _t("report", "pupils")
    assert counts.loc["EN", teacher] == 1
    assert counts.loc["EN", pupils] == 1
    assert counts.loc["L", teacher] == 1 and counts.loc["L", pupils] == 0
    assert counts.loc["I", teacher] == 0 and counts.loc["I", pupils] == 1
    # Gesamtsumme bleibt die Gesamthäufigkeit — die Aufteilung verliert nichts.
    assert int(counts.to_numpy().sum()) == 4


def test_code_counts_by_group_counts_primary_code_only():
    # Gleiche Regel wie Balken/Chip/Zeitverlauf: Spalte 2 zählt nicht mit.
    df = _grouped_df(["LEHRER", "S1"], ["EN (90 %)", "H (80 %)"],
                     ["L (61 %)", "ÄN (30 %)"])
    counts = code_counts_by_group(df, _t, "LEHRER")
    assert sorted(counts.index) == ["EN", "H"]
    assert int(counts.to_numpy().sum()) == 2


def test_code_counts_by_group_resolves_teacher_aliases():
    # Der Fallback-Pfad ohne Transkript trägt die rohen LLM-Sprecherlabels;
    # "Lehrperson" muss dort genauso als Lehrkraft zählen wie "LEHRER".
    df = _grouped_df(["Lehrperson", "S1"], ["EN (90 %)", "I (60 %)"])
    counts = code_counts_by_group(df, _t, "LEHRER")
    assert counts.loc["EN", _t("report", "teacher")] == 1
    assert counts.loc["I", _t("report", "pupils")] == 1


def test_code_counts_by_group_without_teacher_name_is_empty_split():
    # Ohne Lehrkraft-Namen gibt es keine Rollen-Information — dann landet
    # alles in der Schüler-Spalte und der Plot fällt einfarbig zurück.
    df = _grouped_df(["LEHRER", "S1"], ["EN (90 %)", "I (60 %)"])
    counts = code_counts_by_group(df, _t, "")
    assert counts[_t("report", "teacher")].sum() == 0
    assert counts[_t("report", "pupils")].sum() == 2


def test_code_counts_by_group_empty_inputs():
    assert code_counts_by_group(None, _t, "LEHRER").empty
    assert code_counts_by_group(pd.DataFrame(), _t, "LEHRER").empty
    # Turns vorhanden, aber nichts codiert:
    assert code_counts_by_group(_grouped_df(["LEHRER"], [""]), _t, "LEHRER").empty


def test_qual_plot_stacks_by_speaker_group():
    df = _grouped_df(["LEHRER", "S1", "LEHRER"],
                     ["EN (90 %)", "EN (80 %)", "L (70 %)"])
    ax = build_qual_plot(df, _t, "light", "LEHRER")
    # Zwei Container = zwei gestapelte Segmente, Legende benennt beide Gruppen.
    assert len(ax.containers) == 2
    legend_labels = [txt.get_text() for txt in ax.get_legend().get_texts()]
    assert legend_labels == [_t("report", "teacher"), _t("report", "pupils")]


def test_qual_plot_falls_back_to_single_colour_for_one_group():
    # Nur die Lehrkraft codiert: ein Stapel mit toter zweiter Farbe wäre
    # irreführend — dann die klassische einfarbige Darstellung ohne Legende.
    df = _grouped_df(["LEHRER", "LEHRER"], ["EN (90 %)", "L (70 %)"])
    ax = build_qual_plot(df, _t, "light", "LEHRER")
    assert len(ax.containers) == 1
    assert ax.get_legend() is None


def test_docx_code_group_table_totals():
    from docx import Document
    from talktrace_ai.utils.reports import _docx_code_group_table

    df = _grouped_df(["LEHRER", "S1", "LEHRER"],
                     ["EN (90 %)", "EN (80 %)", "L (70 %)"])
    counts = code_counts_by_group(df, _t, "LEHRER")
    doc = Document()
    _docx_code_group_table(doc, counts)
    table = doc.tables[0]
    # Kopfzeile + zwei Codes + Summenzeile; letzte Spalte ist die Zeilensumme.
    # Zeilen über die POSITION prüfen, nicht über lokalisierte Labels: der
    # DOCX-Builder übersetzt gegen die eingestellte App-Sprache, die Tests
    # dürfen nicht von der Config des ausführenden Rechners abhängen.
    assert len(table.rows) == 4
    body = {r.cells[0].text: [c.text for c in r.cells[1:]] for r in table.rows[1:3]}
    assert body["EN"] == ["1", "1", "2"]
    assert body["L"] == ["1", "0", "1"]
    assert [c.text for c in table.rows[-1].cells[1:]] == ["2", "1", "3"]


def test_docx_code_group_table_skipped_without_data():
    from docx import Document
    from talktrace_ai.utils.reports import _docx_code_group_table

    doc = Document()
    _docx_code_group_table(doc, None)
    _docx_code_group_table(doc, pd.DataFrame())
    assert len(doc.tables) == 0


# ---------------------------------------------------------------------------
# Konfidenz-Bänder (Report-Einfärbung)
# ---------------------------------------------------------------------------

def test_extract_confidence_from_cell():
    assert extract_confidence("EN (92 %)") == 92
    assert extract_confidence("EN (7 %)") == 7
    # Handkorrigierte / konfidenzlose Zellen tragen keinen Wert.
    assert extract_confidence("EN") is None
    assert extract_confidence("") is None


@pytest.mark.parametrize("value,expected", [
    (100, "high"), (90, "high"),           # Untergrenze "sicher"
    (89, "medium"), (60, "medium"), (50, "medium"),
    (49, "low"), (0, "low"),               # Obergrenze "sehr unsicher"
    (None, None),
])
def test_confidence_band_thresholds(value, expected):
    assert confidence_band(value) == expected


def test_confidence_thresholds_match_prompt_anchors():
    # Die Schwellen sind bewusst an die Kalibrier-Anker im Prompt gebunden
    # ("90+ NUR bei eindeutiger Passung", "unter 50 = spekulativ"). Laufen
    # sie auseinander, markiert der Report als sicher, was das Modell laut
    # Instruktion nicht als sicher gemeint hat.
    from talktrace_ai.localization.translation import TRANSLATIONS
    assert CONFIDENCE_HIGH_MIN == 90
    assert CONFIDENCE_LOW_MAX == 49  # "unter 50"
    for lang in ("de", "en"):
        txt = TRANSLATIONS[lang]["sidebar"]["prompt_multi_coding_on"]
        assert str(CONFIDENCE_HIGH_MIN) in txt
        assert str(CONFIDENCE_LOW_MAX + 1) in txt


def test_confidence_band_of_cell_roundtrip():
    assert confidence_band_of_cell("EN (92 %)") == "high"
    assert confidence_band_of_cell("L (61 %)") == "medium"
    assert confidence_band_of_cell("ÄN (30 %)") == "low"
    assert confidence_band_of_cell("EN") is None


def test_docx_code_cells_carry_mark_and_shading():
    from docx import Document
    from talktrace_ai.utils.reports import (
        CONFIDENCE_MARKS, _docx_quali_section, translate as _rt,
    )

    # Spalten gegen DIESELBE Übersetzungsquelle bauen, die der DOCX-Builder
    # nutzt (Config-Sprache) — sonst hängt der Test an der App-Sprache des
    # ausführenden Rechners.
    c1, c2 = code_column_names(_rt)
    df = pd.DataFrame({
        "#": [1, 2, 3],
        _rt("report", "speaker"): ["LEHRER", "S1", "LEHRER"],
        _rt("report", "teacher_statement"): ["a", "b", "c"],
        c1: ["EN (92 %)", "H (61 %)", "L"],   # sicher / unsicher / ohne Wert
        c2: ["", "", ""],
    })
    doc = Document()
    # Die Sektion bettet immer den Plot ein — echten mitgeben statt None.
    _docx_quali_section(doc, 3, build_qual_plot(df, _rt, "light", "LEHRER"), df)
    table = doc.tables[-1]
    cells = [r.cells[-2].text for r in table.rows[1:]]  # Spalte "Code 1"
    assert cells[0].endswith(CONFIDENCE_MARKS["high"])
    assert cells[1].endswith(CONFIDENCE_MARKS["medium"])
    # Ohne Konfidenzwert bleibt die Zelle unmarkiert — eine handkorrigierte
    # Zuordnung darf nicht wie eine spekulative Modell-Zuordnung aussehen.
    assert cells[2] == "L"
    # Schattierung nur auf den beiden markierten Zellen.
    shaded = [len(r.cells[-2]._tc.xpath('.//w:shd')) for r in table.rows[1:]]
    assert shaded == [1, 1, 0]


def test_confidence_legend_derives_bounds_from_code():
    from talktrace_ai.utils.reports import _confidence_legend_text

    txt = _confidence_legend_text()
    # Die Legende baut ihre Grenzen aus den Konstanten — kein hartcodierter
    # String, der beim Verschieben der Schwellen stehen bliebe.
    assert f"≥ {CONFIDENCE_HIGH_MIN} %" in txt
    assert f"{CONFIDENCE_LOW_MAX + 1}–{CONFIDENCE_HIGH_MIN - 1} %" in txt
    assert f"< {CONFIDENCE_LOW_MAX + 1} %" in txt
