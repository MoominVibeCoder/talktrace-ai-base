"""Tests für den Rück-Import einer Codierungs-Tabelle aus einem Report.

Der Kern ist ein echter Rundlauf: mit generate_report2 exportieren und mit
parse_report wieder einlesen. Damit hängen Export und Import aneinander —
wer eine Report-Spalte umbenennt, bekommt hier sofort einen roten Test statt
später einen stillen Importfehler beim Nutzer.
"""
import matplotlib
import pandas as pd
import pytest

matplotlib.use("Agg")

from talktrace_ai.localization.translation import TRANSLATIONS
from talktrace_ai.utils.qualitative import build_qual_plot, code_column_names
from talktrace_ai.utils.report_import import (
    ReportImportError,
    coded_turn_count,
    parse_report,
    turns_to_transcript,
)
from talktrace_ai.utils.reports import generate_report2, translate as _rt
from talktrace_ai.utils.stats import dialog_stats

TEACHER = "LEHRER"

# (Sprecher, Code 1, Code 2) — mit Konfidenzen, einem uncodierten Turn und
# einer handkorrigierten Zelle ohne Wert.
ROWS = [
    (TEACHER, "EN (92 %)", "L (61 %)"),
    ("S1",    "N (78 %)",  ""),
    (TEACHER, "L (95 %)",  ""),
    ("S2",    "H (40 %)",  ""),
    ("S1",    "I",         ""),      # handkorrigiert, ohne Konfidenz
    (TEACHER, "",          ""),      # uncodiert
]


def _table(translate=_rt):
    c1, c2 = code_column_names(translate)
    return pd.DataFrame({
        "#": range(1, len(ROWS) + 1),
        translate("report", "speaker"): [r[0] for r in ROWS],
        translate("report", "teacher_statement"): [
            f"Dies ist Äußerung Nummer {i + 1}." for i in range(len(ROWS))
        ],
        c1: [r[1] for r in ROWS],
        c2: [r[2] for r in ROWS],
    })


def _export(path, fmt, table=None):
    table = _table() if table is None else table
    generate_report2(
        str(path), "Testgruppe", 22, 6, 27.3,
        {"num": 3, "words": 60, "mean_sd": "20.0 (2.0)"},
        {"num": 3, "words": 30, "mean_sd": "10.0 (1.0)"},
        None, len(table),
        plot_impulse_coding=build_qual_plot(table, _rt, "light", TEACHER),
        impulse_table=table,
        sections={"quant": False, "quali": True, "legend": False},
        output_format=fmt,
    )
    return path


@pytest.mark.parametrize("fmt,suffix", [
    ("docx", ".docx"), ("xlsx", ".xlsx"), ("html", ".html"), ("csv", ".zip"),
])
def test_roundtrip_export_then_import(tmp_path, fmt, suffix):
    path = _export(tmp_path / f"report{suffix}", fmt)
    frame = parse_report(str(path))
    assert list(frame.columns) == ["Sprecher", "Impuls", "Shortcode"]
    assert len(frame) == len(ROWS)
    # Konfidenz-Suffixe sind weg, der Primärcode bleibt; der Nebencode aus
    # Spalte 2 wird verworfen wie in jeder anderen Häufigkeits-Auswertung.
    assert frame["Shortcode"].tolist() == ["EN", "N", "L", "H", "I", ""]
    assert frame["Sprecher"].tolist() == [r[0] for r in ROWS]
    assert coded_turn_count(frame) == 5   # der uncodierte Turn zählt nicht


def test_import_finds_table_by_headers_not_position(tmp_path):
    # Der DOCX-Report stellt der Turn-Tabelle die Verteilung nach
    # Sprechergruppe voran — ein tables[0] liefe auf die falsche Tabelle.
    from docx import Document
    from talktrace_ai.utils.qualitative import code_counts_by_group

    table = _table()
    counts = code_counts_by_group(table, _rt, TEACHER)
    path = tmp_path / "report.docx"
    generate_report2(
        str(path), "G", 22, 6, 27.3,
        {"num": 3, "words": 60, "mean_sd": "20.0 (2.0)"},
        {"num": 3, "words": 30, "mean_sd": "10.0 (1.0)"},
        None, len(table),
        plot_impulse_coding=build_qual_plot(table, _rt, "light", TEACHER),
        impulse_table=table, sections={"quant": False, "quali": True, "legend": False},
        output_format="docx", code_group_df=counts,
    )
    assert len(Document(str(path)).tables) == 2   # Voraussetzung des Tests
    frame = parse_report(str(path))
    assert frame["Shortcode"].tolist() == ["EN", "N", "L", "H", "I", ""]


def test_import_strips_legacy_confidence_marks(tmp_path):
    # Reports einer früheren Version tragen Konfidenz-Zeichen hinter dem Code.
    # Die liegen in freier Wildbahn und müssen beim Import abfallen, sonst
    # landet "EN ●●●" als eigener Code im Profil.
    table = _table()
    c1 = code_column_names(_rt)[0]
    table[c1] = ["EN (92 %) ●●●", "N (78 %) ●●", "L (95 %) ●●●",
                 "H (40 %) ●", "I", ""]
    path = _export(tmp_path / "legacy.docx", "docx", table)
    frame = parse_report(str(path))
    assert frame["Shortcode"].tolist() == ["EN", "N", "L", "H", "I", ""]


def test_import_reads_the_other_language(tmp_path):
    # Ein deutscher Report muss sich in einer englischen App einlesen lassen
    # und umgekehrt. Das Dokument wird direkt gebaut, nicht über den
    # Report-Builder: der übersetzt gegen die eingestellte App-Sprache und
    # könnte einen fremdsprachigen Report gar nicht erzeugen — genau deshalb
    # muss der IMPORT beide Sprachen kennen.
    from docx import Document

    other = ("en" if _rt("report", "speaker")
             == TRANSLATIONS["de"]["report"]["speaker"] else "de")
    tr = TRANSLATIONS[other]["report"]
    headers = ["#", tr["speaker"], tr["teacher_statement"],
               f"{tr['shortcode']} 1", f"{tr['shortcode']} 2"]

    doc = Document()
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for i, (speaker, c1, c2) in enumerate(ROWS, start=1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = speaker
        cells[2].text = f"Statement number {i}."
        cells[3].text = c1
        cells[4].text = c2
    path = tmp_path / "other.docx"
    doc.save(str(path))

    frame = parse_report(str(path))
    assert frame["Shortcode"].tolist() == ["EN", "N", "L", "H", "I", ""]
    assert frame["Sprecher"].tolist() == [r[0] for r in ROWS]


def test_metrics_reconstruct_from_imported_report(tmp_path):
    # Der Import trägt sich selbst: aus den Turns wird ein Transkript-Text
    # zusammengesetzt, auf dem die normalen Statistikfunktionen laufen —
    # kein zweiter Rechenpfad neben dem Analyse-Lauf.
    path = _export(tmp_path / "report.docx", "docx")
    frame = parse_report(str(path))
    transcript = turns_to_transcript(frame, TEACHER)
    assert transcript.count("\n") + 1 == len(ROWS)
    stats = dialog_stats(transcript, TEACHER)
    teacher_row = stats[stats["Sprecher"] == TEACHER].iloc[0]
    assert teacher_row["Anzahl_Beitraege"] == 3     # drei Lehrkraft-Turns
    students = stats[stats["Sprecher"] != TEACHER].iloc[0]
    assert students["Anzahl_Beitraege"] == 3


def test_import_rejects_unsupported_format(tmp_path):
    path = tmp_path / "report.pdf"
    path.write_bytes(b"%PDF-1.4 not really a pdf")
    with pytest.raises(ReportImportError) as exc:
        parse_report(str(path))
    # Die Nachricht ist ein Localization-Key, den der Handler übersetzt.
    assert str(exc.value) == "import_unsupported_format"
    for lang in ("de", "en"):
        assert TRANSLATIONS[lang]["feedback"][str(exc.value)].strip()


def test_import_without_coding_table_raises(tmp_path):
    # Ein Report ohne qualitative Sektion trägt keine Codierungs-Tabelle.
    from docx import Document

    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "Irgendwas"
    tbl.rows[0].cells[1].text = "Anderes"
    path = tmp_path / "fremd.docx"
    doc.save(str(path))
    with pytest.raises(ReportImportError) as exc:
        parse_report(str(path))
    assert str(exc.value) == "import_no_table"


def test_import_of_broken_file_raises_readable_error(tmp_path):
    path = tmp_path / "kaputt.docx"
    path.write_bytes(b"das ist kein docx")
    with pytest.raises(ReportImportError) as exc:
        parse_report(str(path))
    assert str(exc.value) == "import_unreadable"


def test_suffix_comes_from_the_original_filename(tmp_path):
    # Shiny legt Uploads unter einem Temp-Pfad OHNE Endung ab; das Format
    # muss deshalb aus dem mitgelieferten Dateinamen kommen.
    path = _export(tmp_path / "report.docx", "docx")
    blind = tmp_path / "0"
    blind.write_bytes(path.read_bytes())
    with pytest.raises(ReportImportError):
        parse_report(str(blind))
    frame = parse_report(str(blind), "Mein Report.docx")
    assert len(frame) == len(ROWS)


def test_turns_to_transcript_handles_empty():
    assert turns_to_transcript(None) == ""
    assert turns_to_transcript(pd.DataFrame()) == ""
    assert coded_turn_count(None) == 0


def test_detect_teacher_label_from_report(tmp_path):
    # Ohne diese Ableitung rechnen die Kennzahlen still falsch: heißt die
    # Lehrkraft im Report "L", die App steht aber auf "LEHRER", zählte
    # dialog_stats ALLE Beiträge als Schülerbeiträge.
    from talktrace_ai.utils.report_import import detect_teacher_label

    frame = pd.DataFrame({
        "Sprecher": ["L", "S1", "L", "S2", "L"],
        "Impuls": [f"Text {i}" for i in range(5)],
        "Shortcode": ["EN", "N", "L", "H", ""],
    })
    # Die App-Einstellung kommt im Report nicht vor -> Label aus den Daten.
    assert detect_teacher_label(frame, "LEHRER") == "L"
    # Kommt sie vor, gewinnt sie (der Nutzer weiß, wie sein Material heißt).
    assert detect_teacher_label(frame, "l") == "L"
    # Nur Schülerlabels -> nichts ableitbar.
    students = pd.DataFrame({"Sprecher": ["S1", "S2"], "Impuls": ["a", "b"],
                             "Shortcode": ["", ""]})
    assert detect_teacher_label(students, "LEHRER") is None
    assert detect_teacher_label(None, "LEHRER") is None


def test_teacher_label_drives_the_metrics_split(tmp_path):
    # Regressionsschutz für den stillen Fehler: mit dem richtigen Label
    # werden 3 Lehrkraft- und 2 Schülerbeiträge gezählt. Mit einem falschen
    # Label fallen die Lehrkraft-Turns GANZ aus der Statistik — der Regex in
    # dialog_stats kennt nur <teacher>|S<n>, und "L:" ist dann keines von
    # beidem. Der Report sähe aus, als hätte die Lehrkraft geschwiegen.
    from talktrace_ai.utils.report_import import detect_teacher_label

    frame = pd.DataFrame({
        "Sprecher": ["L", "S1", "L", "S2", "L"],
        "Impuls": [f"Ein Beitrag Nummer {i}" for i in range(5)],
        "Shortcode": ["EN", "N", "L", "H", ""],
    })
    transcript = turns_to_transcript(frame)
    good = dialog_stats(transcript, detect_teacher_label(frame, "LEHRER"))
    assert good[good["Sprecher"] == "L"].iloc[0]["Anzahl_Beitraege"] == 3
    assert good[good["Sprecher"] != "L"].iloc[0]["Anzahl_Beitraege"] == 2

    bad = dialog_stats(transcript, "LEHRER")
    assert bad["Anzahl_Beitraege"].sum() == 2      # die 3 "L:"-Turns fehlen
    assert "L" not in bad["Sprecher"].tolist()
